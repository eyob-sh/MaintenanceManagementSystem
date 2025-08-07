from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.hashers import make_password
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.http import Http404

from django.db.models import Count, Q, Sum,F
from datetime import datetime, timedelta
from django.db.models.functions import TruncMonth
from django.utils import timezone
from django.http import JsonResponse
from django.urls import reverse
from django.db.models.deletion import ProtectedError, Collector
from django.views.decorators.csrf import csrf_exempt
from django.http import StreamingHttpResponse
from asgiref.sync import sync_to_async
import asyncio
from django.dispatch import receiver
from .signals import notification_created
from django.db import models
from django.core.paginator import Paginator
from easyaudit.models import CRUDEvent, LoginEvent, RequestEvent

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, KeepTogether, PageBreak
from reportlab.lib.units import inch
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.styles import ParagraphStyle
import tempfile
import os
from django.core.files.uploadedfile import InMemoryUploadedFile
from django.http import StreamingHttpResponse
import time
from decimal import Decimal, InvalidOperation
from django.db.models import Case, When, Value, IntegerField
from django.db.models.functions import Coalesce
from django.db import close_old_connections


import io

import json
from import_export import resources
from import_export.formats import base_formats
from import_export.admin import ImportMixin, ExportMixin
from .models import MaintenanceRecord, MaintenanceTask, Manufacturer, SparePart, SparePartUsage, RestockSparePart, DecommissionedEquipment,Equipment, Notification,WorkOrder,Branch,UserProfile, Task, TaskGroup, TaskCompletion, SparePartRequest, SparePartTransaction, TechnicianSparePart

from .forms import EquipmentForm, SparePartForm, MaintenanceRecordForm, ManufacturerForm, WorkOrderForm, SparePartUsageForm, DecommissionedEquipmentForm, MaintenanceTaskForm,  RestockSparePartForm, BranchForm, UserProfileForm, TaskForm, TaskGroupForm
from .resources import EquipmentResource
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from django.contrib.auth.decorators import user_passes_test
from django.views.decorators.cache import never_cache


def is_md(user):
    return user.is_authenticated and hasattr(user, 'userprofile') and user.userprofile.role == 'MD manager'

def is_tec(user):
    return user.is_authenticated and hasattr(user, 'userprofile') and user.userprofile.role == 'TEC'

def is_mo(user):
    return user.is_authenticated and hasattr(user, 'userprofile') and user.userprofile.role == 'MO'

def is_im(user):
    return user.is_authenticated and hasattr(user, 'userprofile') and user.userprofile.role == 'IM'

def is_cl(user):
    return user.is_authenticated and hasattr(user, 'userprofile') and user.userprofile.role == 'CL'

def is_ad(user):
    return user.is_authenticated and hasattr(user, 'userprofile') and user.userprofile.role == 'AD'


# Create your views here.


DEPARTMENT_CHOICES = [
    ('MD', 'Maintenance Department'),
    ('CD', 'Chemicals Department'),
    
]

# Define choices for roles
ROLE_CHOICES = [
    ('MD manager', 'Maintenance Department Manager'),
    ('TEC', 'Technician'),
    ('MO', 'Maintenance Oversight'),
    ('IM', 'Inventory Manager'),
    ('CL', 'Client'),
    ('AD', 'Admin'),
    
]

def custom_404_view(request, exception):
    return render(request, '404.html', status=404)

def custom_403_view(request, exception):
    return render(request, '404.html', status=403)

@never_cache
def loginPage(request):

    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')

        # try:
        #     user = user.objects.get(username = username)
        
        # except: 
        #     messages.error(request, 'user does not exist')


        user = authenticate(request, username = username, password = password)

        if user is not None and not user.is_superuser:
            login(request, user)
            if (request.user.userprofile.role in 'MD manager, TEC'):
                return redirect('maintenance_dashboard')
            elif (request.user.userprofile.role in 'MO'):
                return redirect('maintenance_oversight_dashboard')
            elif (request.user.userprofile.role in 'AD'):
                return redirect('dashboard')
            elif(request.user.userprofile.role in 'CL'):
                return redirect('client_dashboard')
            elif(request.user.userprofile.role in 'IM'):
                return redirect('inventory_dashboard')
        else: 
            messages.error(request, 'Incorrect Username or Password')
        
    context = {}
    return render(request, 'login_page.html', context)

@never_cache
def logoutUser(request):
    logout(request)
    return redirect('login')




@login_required
def my_profile(request):
    # Get the current user's profile
    user_profile = UserProfile.objects.get(user=request.user)
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    
    if request.method == 'POST':
        # Handle email update
        new_email = request.POST.get('email').strip()
        if new_email != request.user.email:
            # Update the user's email
            request.user.email = new_email if new_email else ''  # Allow empty email
            request.user.save()
            messages.success(request, 'Email updated successfully!')

        # Handle password update
        password = request.POST.get('password')
        confirm_password = request.POST.get('confirmPassword')
        if password and confirm_password:
            if password == confirm_password:
                # Update the user's password
                request.user.password = make_password(password)
                request.user.save()
                messages.success(request, 'Password updated successfully! Please log in again with your new password.')
            else:
                messages.error(request, 'Passwords do not match!')
                return redirect('my_profile')
        
        # messages.success(request, 'Profile updated successfully!')
        return redirect('my_profile')
    
    return render(request, 'my_profile.html', {
        'user_profile': user_profile,
                'active_page': 'profile',
                'notifications':notifications,
                'latest_notification_id': latest_notification.id if latest_notification else 0,

        
    })

@user_passes_test(lambda u: is_ad(u), login_url=None)
def user_profile_list(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    if request.user.userprofile.branch.name == "Head Office":
        user_profiles = UserProfile.objects.all()
    else:
        user_profiles = UserProfile.objects.filter(branch = request.user.userprofile.branch)
    context = {
        'active_page': 'user_profile_list',
        'title': 'User Profiles',
        'item_list': user_profiles,
        'edit_url': 'edit_user_profile',
        'notifications':notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'user_profile_list.html', context)

@user_passes_test(lambda u: is_ad(u), login_url=None)
def add_user_profile_page(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    # Fetch all branches from the database
    branches = Branch.objects.all()

    # Debugging: Print the choices to the console
    print("Roles:", UserProfile.ROLE_CHOICES)
    print("Branches:", branches)

    # Pass the DEPARTMENT_CHOICES, ROLE_CHOICES, and branches to the template
    context = {
        'roles': UserProfile.ROLE_CHOICES,
        'branches': branches,  # Pass the branches queryset
         'active_page': 'user_profile_list',
         'notifications':notifications,
         'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'add_user_profile.html', context)

@user_passes_test(lambda u: is_ad(u), login_url=None)
def add_user_profile(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    if request.method == 'POST':
        first_name = request.POST.get('firstName')
        last_name = request.POST.get('lastName')
        username = request.POST.get('username')
        email = request.POST.get('email')  # Email is optional
        password = request.POST.get('password')
        role = request.POST.get('role')  # Role as a string
        branch_id = request.POST.get('branch')
        is_active = request.POST.get('is_active')  # Default to active
        is_active = is_active.capitalize()

        # Check if the username already exists
        if User.objects.filter(username=username).exists():
            messages.error(request, 'Username already exists')
            return redirect('add_user_profile_page')

        # Create the User
        user = User.objects.create_user(
            username=username,
            email=email,  # Email is optional
            password=password,
            first_name=first_name,
            last_name=last_name,
            is_active=is_active  # Set user active status
        )

        # Create UserProfile
        user_profile = UserProfile(
            user=user,
            branch_id=branch_id,
            role=role  # Pass the role string
        )
        user_profile.save()

        messages.success(request, 'User created successfully')
        return redirect('user_profile_list')  # Redirect to the user list page

    return render(request, 'add_user_profile.html', { 'active_page': 'user_profile_list','latest_notification_id': latest_notification.id if latest_notification else 0,})
def edit_user_profile(request, id):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    user_profile = get_object_or_404(UserProfile, id=id)
    
    if request.method == 'POST':
        form = UserProfileForm(request.POST, instance=user_profile)
        if form.is_valid():
            # Update the user's active status
            is_active = request.POST.get('is_active', 'on') == 'on'
            user_profile.user.is_active = is_active
            user_profile.user.save()  # Save the User model

            # Handle password update
            password = request.POST.get('password')
            confirm_password = request.POST.get('confirmPassword')
            if password and confirm_password:
                if password == confirm_password:
                    # Update the user's password
                    user_profile.user.password = make_password(password)
                    user_profile.user.save()
                    messages.success(request, 'Password updated successfully!')
                else:
                    messages.error(request, 'Passwords do not match!')
                    return redirect('edit_user_profile', id=id, )
            

            form.save()  # Save the UserProfile model
            messages.success(request, 'User Profile updated successfully!')
            return redirect('user_profile_list')
    else:
        form = UserProfileForm(instance=user_profile)
    
    return render(request, 'edit_user_profile.html', {
        'form': form,
        'user_profile': user_profile,
         'active_page': 'user_profile_list',
         'notifications':notifications,
         'latest_notification_id': latest_notification.id if latest_notification else 0,
    })
def check_username(request):
    username = request.GET.get('username')
    exists = User.objects.filter(username=username).exists()
    return JsonResponse({'exists': exists})

@user_passes_test(lambda u: is_ad(u), login_url=None)
def branch_list(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    branches = Branch.objects.all()
    context = {
        'active_page': 'branch_list',
        'title': 'Branches',
        'item_list': branches,
        'edit_url': 'edit_branch',
        'notifications':notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'branch_list.html', context)

@user_passes_test(lambda u: is_ad(u), login_url=None)
def add_branch_page(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    return render(request, 'add_branch.html',{ 'active_page': 'branch_list','notifications':notifications,'latest_notification_id': latest_notification.id if latest_notification else 0,})



@user_passes_test(lambda u: is_ad(u), login_url=None)
def add_branch(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    if request.method == 'POST':
        form = BranchForm(request.POST)  # Bind the form to the POST data
        if form.is_valid():  # Validate the form
            form.save()  # Save the form data to the database
            messages.success(request, 'Branch added successfully!')
            return redirect('branch_list')  # Redirect to the branch list page
    else:
        form = BranchForm()  # Create an empty form for GET requests

    # Render the form in the template
    return render(request, 'add_branch.html', {'form': form,  'active_page': 'branch_list','notifications':notifications,'latest_notification_id': latest_notification.id if latest_notification else 0,})

@user_passes_test(lambda u: is_ad(u), login_url=None)
def edit_branch(request, id):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    branch = get_object_or_404(Branch, id=id)
    if request.method == 'POST':
        form = BranchForm(request.POST, instance=branch)
        if form.is_valid():
            form.save()
            messages.success(request, 'Branch updated successfully!')
            return redirect('branch_list')
    else:
        form = BranchForm(instance=branch)
    return render(request, 'edit_branch.html', {'form': form, 'branch': branch, 'active_page': 'branch_list','notifications':notifications,'latest_notification_id': latest_notification.id if latest_notification else 0,})

def get_notifications(user):
    return Notification.objects.filter(user=user, is_read=False).order_by('-timestamp')[:5]


@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_mo(u), login_url=None)
def manufacturer_list(request):
    user_branch = request.user.userprofile.branch
    
    
    if request.user.userprofile.role in ['MO', 'Maintenance Oversight']:
         manufacturers = Manufacturer.objects.all()  
    else:
        manufacturers = Manufacturer.objects.filter(site = user_branch)
        
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    context = {
        'active_page': 'manufacturer_list',
        'title': 'Manufacturers',
        'item_list': manufacturers,
        'edit_url': 'edit_manufacturer',
        'delete_url': 'delete_manufacturer',
        'notifications':notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'manufacturer_list.html', context)

@user_passes_test(lambda u: is_tec(u) or is_md(u) , login_url=None)
def add_manufacturer_page(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    
    context = {
    'active_page': 'manufacturer_list', 'notifications':notifications,'latest_notification_id': latest_notification.id if latest_notification else 0,}
    
    return render(request, 'add_manufacturer.html', context)

@user_passes_test(lambda u: is_tec(u) or is_md(u), login_url=None)
def add_manufacturer(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    # Get the user's branch
    if request.user.is_authenticated:
        branch = request.user.userprofile.branch  # Ensure this matches your UserProfile model
    else:
        branch = None  # Handle unauthenticated users

    if request.method == 'POST':
        if request.user.userprofile.role in ["MD manager", "TEC"]:
            # Get form data
            name = request.POST.get('name')
            description = request.POST.get('description')
            site = request.POST.get('site')  # This will be the branch ID
            contact_email = request.POST.get('contact_email')
            contact_phone_number = request.POST.get('contact_phone_number')
            address = request.POST.get('address')

            # Validate required fields
            if not name or not site:
                messages.error(request, 'Please fill out all required fields.')
                return render(request, 'add_manufacturer.html', {'active_page': 'manufacturer_list', 'branch': branch, 'notifications':notifications})

            try:
                # Create the manufacturer
                Manufacturer.objects.create(
                    name=name,
                    description=description,
                    site_id=site,  # Use site_id to assign the branch
                    contact_email=contact_email,
                    contact_phone_number=contact_phone_number,
                    address=address
                )
                messages.success(request, 'Manufacturer added successfully!')
                return redirect('manufacturer_list')
            except Exception as e:
                messages.error(request, f'An error occurred: {str(e)}')
                return render(request, 'add_manufacturer.html', {'active_page': 'manufacturer_list', 'branch': branch, 'notifications':notifications})
        else:
            messages.error(request, f'You do not have the permsission to perform this task')

    return render(request, 'add_manufacturer.html', {'active_page': 'manufacturer_list', 'branch': branch, 'notifications':notifications,'latest_notification_id': latest_notification.id if latest_notification else 0,})

@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_mo(u), login_url=None)
def edit_manufacturer(request, id):
    manufacturer = get_object_or_404(Manufacturer, id=id)
    branch = request.user.userprofile.branch
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()
    
    # Initialize the form variable
    form = ManufacturerForm(instance=manufacturer)

    if request.method == 'POST':
        if request.user.userprofile.role in ["MD manager", "TEC"]:
            form = ManufacturerForm(request.POST, instance=manufacturer)
            if form.is_valid():
                form.save()
                messages.success(request, 'Manufacturer updated successfully!')
                return redirect('manufacturer_list')
        else:
            messages.error(request, 'You do not have the permission to perform this task')

    return render(request, 'edit_manufacturer.html', {
        'form': form,
        'manufacturer': manufacturer,
        'active_page': 'manufacturer_list',
        'branch': branch,
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    })

@user_passes_test(lambda u: is_tec(u) or is_md(u), login_url=None)
def delete_manufacturer(request, id):
    manufacturer = get_object_or_404(Manufacturer, id=id)
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    
    if request.method == 'POST':
        if request.user.userprofile.role in ["MD manager", "TEC"]:
            try:
                manufacturer.delete()
                messages.success(request, 'The manufacturer was deleted successfully.')
                return redirect(reverse('manufacturer_list'))
            except ProtectedError as e:
                messages.error(request, 'This equipment cannot be deleted because it is referenced by other objects.')
                return render(request, 'confirm_delete_protected.html', {
                    'object': manufacturer,
                    'protected_objects': e.protected_objects,
                    'model_name': 'Manufacturer',  # Dynamic model name
                    'active_page':'manufacturer_list',
                    'notifications':notifications,
                    'latest_notification_id': latest_notification.id if latest_notification else 0,
                })
        else:
            messages.error(request, f'You do not have the permsission to perform this task')
    
    return render(request, 'confirm_delete.html', {
        'object': manufacturer,
        'model_name': 'Manufacturer',  # Dynamic model name
        'active_page':'manufacturer_list',

        'notifications':notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,

    })


#----------------------------------------------------------------------------------
@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_mo(u) or is_cl(u), login_url=None)
def work_order_list(request):
    user_branch = request.user.userprofile.branch
    
    if request.user.userprofile.role in ['MO', 'Maintenance Oversight']:
         work_orders = WorkOrder.objects.all()  # Show all equipment for MO
    elif request.user.userprofile.role in ['MD manager', 'TEC']:
        work_orders = WorkOrder.objects.filter(branch = user_branch)
    elif request.user.userprofile.role in ['CL']:
        work_orders = WorkOrder.objects.filter(branch = user_branch, requester = request.user)

   
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    context = {
        'active_page': 'work_order_list',
        'title': 'Work Orders',
        'item_list': work_orders,
        'edit_url': 'edit_work_order',
        'notifications':notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'work_order_list.html', context)
@user_passes_test(lambda u: is_md(u) or is_cl(u),  login_url=None)
def add_work_order_page(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()
    user_branch = request.user.userprofile.branch

    equipments = Equipment.objects.filter(branch=user_branch, decommissioned = False)
    
    return render(request, 'add_work_order.html', {
        'equipments': equipments,
        'branch': user_branch,
        'notifications': notifications,
        'active_page': 'work_order_list',
        'latest_notification_id': latest_notification.id if latest_notification else 0,

    })
@user_passes_test(lambda u: is_md(u) or is_cl(u),  login_url=None)
def add_work_order(request):
    notifications = get_notifications(request.user)

    if request.method == 'POST':
        branch = request.POST.get('branch')
        equipment_id = request.POST.get('equipment')
        location = request.POST.get('location')
        description = request.POST.get('description')
        remark = request.POST.get('remark')
        status = 'Pending'  # Default status
        if(request.user.userprofile.role == 'MD manager'):
            status = 'Price_Confirmed'

        if not branch or not description:
            messages.error(request, 'Please fill out all required fields.')
            return redirect('add_work_order_page')

        try:
            work_order = WorkOrder.objects.create(
                requester=request.user,
                branch_id=branch,
                equipment_id=equipment_id,
                location = location,
                description=description,
                remark = remark,
                status=status
            )
            messages.success(request, 'Work Order added successfully!')

            # Send notification to the MD Manager of the same branch
            if request.user.userprofile.role == 'CL':
                managers = User.objects.filter(userprofile__branch=branch, userprofile__role='MD manager')

                # Iterate over the managers and create a notification for each one
                for manager in managers:
                    Notification.objects.create(
                        user=manager,
                        type="work_order",
                        message=f'New work order: {work_order.location}.',
                    )
                    notification_created.send(sender=Notification)
            
            return redirect('work_order_list')
        except Exception as e:
            messages.error(request, f'An error occurred: {str(e)}')
            return redirect('add_work_order_page')

    return redirect('add_work_order_page')

@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_cl(u) or is_mo(u),  login_url=None)
def edit_work_order(request, id):
    user_branch = request.user.userprofile.branch
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()
    spare_parts = TechnicianSparePart.objects.filter(
        technician=request.user
    ).select_related('spare_part')
    work_order = get_object_or_404(WorkOrder, id=id)
    assigned_technician_ids = work_order.assigned_technicians.values_list('id', flat=True)
    equipments = Equipment.objects.filter(branch=user_branch)

    if request.method == 'POST':
        if request.user.userprofile.role in ["MD manager", "TEC"]:
            # Get form data
            equipment_id = request.POST.get('equipment')
            remark = request.POST.get('remark')
            assigned_technicians = request.POST.getlist('assigned_technicians[]')
            spare_parts_post = request.POST.getlist('spare_parts[]')
            spare_part_quantities = request.POST.getlist('spare_part_quantities[]')

            try:
                # Step 1: Update the equipment
                if equipment_id:
                    equipment = Equipment.objects.get(id=equipment_id)
                    work_order.equipment = equipment
                work_order.remark = remark
                work_order.save()

                #include price estimation 

                if request.user.userprofile.role == 'MD manager':
                    # Convert assigned_technicians to a set of integers
                    new_technician_ids = set(map(int, assigned_technicians))
                    current_technician_ids = set(assigned_technician_ids)

                    # Check if there are changes in assigned technicians
                    if new_technician_ids != current_technician_ids:
                        work_order.assigned_technicians.set(assigned_technicians)

                        # Notify only newly assigned technicians
                        newly_assigned_technicians = new_technician_ids - current_technician_ids
                        for technician_id in newly_assigned_technicians:
                            technician = User.objects.get(id=technician_id)
                            Notification.objects.create(
                                user=technician,
                                type="work_order",
                                message=f'You have been assigned a price estimation task for the work order : {work_order}.',
                            )
                            notification_created.send(sender=Notification)



                # Step 2: Update assigned technicians (only for MD manager)
                if request.user.userprofile.role == 'MD manager' and work_order.status == 'Price_Confirmed':
                    # Convert assigned_technicians to a set of integers
                    new_technician_ids = set(map(int, assigned_technicians))
                    current_technician_ids = set(assigned_technician_ids)

                    # Check if there are changes in assigned technicians
                    if new_technician_ids != current_technician_ids:
                        work_order.assigned_technicians.set(assigned_technicians)

                        # Notify only newly assigned technicians
                        newly_assigned_technicians = new_technician_ids - current_technician_ids
                        for technician_id in newly_assigned_technicians:
                            technician = User.objects.get(id=technician_id)
                            Notification.objects.create(
                                user=technician,
                                type="work_order",
                                message=f'You have been assigned a new work order task: {work_order}.',
                            )
                            notification_created.send(sender=Notification)

                # Step 3: Add back the old quantities to the spare parts
                spare_part_usages = SparePartUsage.objects.filter(work_order=work_order)
                for usage in spare_part_usages:
                    tech_spare = TechnicianSparePart.objects.get(
                        technician=request.user,
                        spare_part=usage.spare_part
                    )
                    # spare_part = usage.spare_part
                    tech_spare.used_quantity -= usage.quantity_used
                    tech_spare.save()


                # Step 4: Process the new spare parts and quantities (only for Technician and if not approved)
                if request.user.userprofile.role == 'TEC' and work_order.status != 'Approved':
                    for spare_part_id, quantity_used in zip(spare_parts_post, spare_part_quantities):
                        # Skip if spare_part_id or quantity_used is empty
                        if not spare_part_id.strip() or not quantity_used.strip():
                            continue  # Skip empty fields

                        # Skip if spare_part_id is not a valid integer
                        try:
                            spare_part_id = int(spare_part_id)
                        except ValueError:
                            continue  # Skip invalid spare_part_id (e.g., non-numeric values)

                        # Skip if quantity_used is not a valid integer
                        try:
                            quantity_used = int(quantity_used)
                        except ValueError:
                            continue  # Skip invalid quantities (e.g., non-numeric values)

                        # Get the spare part
                        spare_part = SparePart.objects.get(id=spare_part_id)

                        tech_spare = TechnicianSparePart.objects.get(
                        technician=request.user,
                        spare_part=spare_part
                    )

                        # Check if the new quantity exceeds the available stock
                        if tech_spare.available_quantity < quantity_used:
                            messages.error(request, f'Not enough quantity for {spare_part.name}. Available: {spare_part.quantity}')
                            # Rollback the old quantities
                            for usage in spare_part_usages:
                                tech_spare.used_quantity += usage.quantity_used
                                tech_spare.save()
                            return redirect('edit_work_order', id=work_order.id)

                        # Deduct the new quantity from the spare part
                        # spare_part.quantity -= quantity_used
                        # spare_part.save()
                        # check_low_spare_parts(spare_part)

                        # Create or update the SparePartUsage record
                        SparePartUsage.objects.update_or_create(
                            work_order=work_order,
                            spare_part=spare_part,
                            defaults={'quantity_used': quantity_used},
                        )

                        req = tech_spare.request
                        SparePartTransaction.objects.update_or_create(
                        request = req, 
                        transaction_type='Usage',
                        user=request.user,
                        quantity=quantity_used,
                        notes=f'Used ({quantity_used}) in work order #{work_order}'
                        )

                    # Step 5: Delete any remaining spare part usages that were not in the form
                    SparePartUsage.objects.filter(work_order=work_order).exclude(
                        spare_part_id__in=[int(id) for id in spare_parts_post if id.strip()]
                    ).delete()
                if request.user.userprofile.role == 'TEC':
                    if work_order.status == 'Accepted' or work_order.status == 'Price_Estimated':
                        estimate_price(request, work_order.id)
                

                messages.success(request, 'Work order updated successfully!')
                return redirect('work_order_list')
            except Exception as e:
                messages.error(request, f'An error occurred: {str(e)}')
                return redirect('edit_work_order', id=work_order.id)
        else:
            messages.error(request, f'You do not have the permsission to perform this task')
    # For GET requests, pre-fill the form and spare parts
    context = {
        'work_order': work_order,
        'technicians': User.objects.filter(userprofile__branch=user_branch, userprofile__role='TEC'),
        'assigned_technician_ids': assigned_technician_ids,
        'active_page': 'work_order_list',
        'notifications': notifications,
        'spare_parts': spare_parts,
        'equipments': equipments,
        'selected_equipment_id': work_order.equipment.id if work_order.equipment else None,
        'spare_part_usages': SparePartUsage.objects.filter(work_order=work_order),
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'edit_work_order.html', context)
#------------------------------------------------------------------------------------


@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_mo(u) or is_im(u), login_url=None)
def spare_part_usage_list(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()
    user_branch = request.user.userprofile.branch
    
    if request.user.userprofile.role in ['MO', 'Maintenance Oversight']:
         spare_part_usages = SparePartUsage.objects.all()
    else:
        spare_part_usages = SparePartUsage.objects.filter(
    Q(maintenance_record__branch=user_branch) | Q(work_order__branch=user_branch)
     )
    context = {
        'active_page': 'spare_part_usage_list',
        'title': 'Spare Part Usages',
        'item_list': spare_part_usages,
        'edit_url': 'edit_spare_part_usage',
        'notifications':notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
        
    }
    return render(request, 'spare_part_usage_list.html', context)

@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_mo(u) or is_im(u), login_url=None)
def add_spare_part_usage_page(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()
    user_branch = request.user.userprofile.branch

    maintenance_records = MaintenanceRecord.objects.filter(branch = user_branch)
    spare_parts = SparePart.objects.all()
    return render(request, 'add_spare_part_usage.html', {
        'maintenance_records': maintenance_records,
        'spare_parts': spare_parts,
    'active_page': 'spare_part_usage_list',
    'notifications': notifications,
    'latest_notification_id': latest_notification.id if latest_notification else 0,
    })

@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_mo(u) or is_im(u), login_url=None)
def add_spare_part_usage(request):
    if request.method == 'POST':
        maintenance_record_id = request.POST.get('maintenance_record')
        spare_part_id = request.POST.get('spare_part')
        quantity_used = request.POST.get('quantity_used')

        if not maintenance_record_id or not spare_part_id or not quantity_used:
            messages.error(request, 'Please fill out all required fields.')
            return redirect('add_spare_part_usage_page')

        try:
            SparePartUsage.objects.create(
                maintenance_record_id=maintenance_record_id,
                spare_part_id=spare_part_id,
                quantity_used=quantity_used
            )
            messages.success(request, 'Spare Part Usage added successfully!')
            return redirect('spare_part_usage_list')
        except Exception as e:
            messages.error(request, f'An error occurred: {str(e)}')
            return redirect('add_spare_part_usage_page')

    return redirect('add_spare_part_usage_page')

@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_mo(u) or is_im(u), login_url=None)
def edit_spare_part_usage(request, id):
    spare_part_usage = get_object_or_404(SparePartUsage, id=id)
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()


    if request.method == 'POST':
        if request.user.userprofile.role in ["MD manager", "TEC"]:
            form = SparePartUsageForm(request.POST, instance=spare_part_usage)
            if form.is_valid():
                form.save()
                messages.success(request, 'Spare Part Usage updated successfully!')
                return redirect('spare_part_usage_list')
        else:
            messages.error(request, f'You do not have the permsission to perform this task')
    else:
        form = SparePartUsageForm(instance=spare_part_usage)

    return render(request, 'edit_spare_part_usage.html', {'form': form, 'spare_part_usage': spare_part_usage, 'active_page': 'spare_part_usage_list', 'notifications':notifications,'latest_notification_id': latest_notification.id if latest_notification else 0,})

#------------------------------------------------------------------------------------
@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_mo(u), login_url=None)
def decommissioned_equipment_list(request):
    user_branch = request.user.userprofile.branch
    if request.user.userprofile.role in ['MO', 'Maintenance Oversight']:
        decommissioned_equipments = DecommissionedEquipment.objects.all()  # Show all equipment for MO
    else:
        decommissioned_equipments = DecommissionedEquipment.objects.filter(equipment__branch=user_branch)    
    
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    context = {
        'active_page': 'decommissioned_equipment_list',
        'title': 'Decommissioned Equipments',
        'item_list': decommissioned_equipments,
        'edit_url': 'edit_decommissioned_equipment',
        'notifications':notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'decommissioned_equipment_list.html', context)

@user_passes_test(lambda u: is_tec(u) or is_md(u) , login_url=None)
def add_decommissioned_equipment_page(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()
    equipments = Equipment.objects.filter(branch = request.user.userprofile.branch, decommissioned = False)
    return render(request, 'add_decommissioned_equipment.html', {
        'equipments': equipments,'active_page': 'decommissioned_equipment_list','notifications':notifications,'latest_notification_id': latest_notification.id if latest_notification else 0,
    })

@user_passes_test(lambda u: is_tec(u) or is_md(u) , login_url=None)
def add_decommissioned_equipment(request):
    if request.method == 'POST':
        equipment_id = request.POST.get('equipment')
        decommission_reason = request.POST.get('decommission_reason')
        decommission_date = request.POST.get('decommission_date')

        if not equipment_id or not decommission_reason or not decommission_date:
            messages.error(request, 'Please fill out all required fields.')
            return redirect('add_decommissioned_equipment_page')

        try:
            equipment = Equipment.objects.get(id = equipment_id)
            
            DecommissionedEquipment.objects.create(
                equipment_id=equipment_id,
                decommission_reason=decommission_reason,
                decommission_date=decommission_date
            )
            equipment.decommissioned = True
            
            
            messages.success(request, 'Decommissioned Equipment added successfully!')
            return redirect('decommissioned_equipment_list')
        except Exception as e:
            messages.error(request, f'An error occurred: {str(e)}')
            return redirect('add_decommissioned_equipment_page')

    return redirect('add_decommissioned_equipment_page')

@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_mo(u) , login_url=None)
def edit_decommissioned_equipment(request, id):
    decommissioned_equipment = get_object_or_404(DecommissionedEquipment, id=id)
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()


    if request.method == 'POST':
        if request.user.userprofile.role in ["MD manager", "TEC"]:
            form = DecommissionedEquipmentForm(request.POST, instance=decommissioned_equipment)
            if form.is_valid():
                form.save()
                messages.success(request, 'Decommissioned Equipment updated successfully!')
                return redirect('decommissioned_equipment_list')
        else:
            messages.error(request, f'You do not have the permsission to perform this task')
        
    else:
        form = DecommissionedEquipmentForm(instance=decommissioned_equipment)

    return render(request, 'edit_decommissioned_equipment.html', {'form': form, 'decommissioned_equipment': decommissioned_equipment, 'active_page': 'decommissioned_equipment_list','notifications':notifications,'latest_notification_id': latest_notification.id if latest_notification else 0,})


#--------------------------------------------------------------------------------------
@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_mo(u), login_url=None)
def maintenance_task_list(request):
    user_branch = request.user.userprofile.branch
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()


    maintenance_tasks = MaintenanceTask.objects.all()
    context = {
        'active_page': 'maintenance_task_list',
        'title': 'Maintenance Task',
        'item_list': maintenance_tasks,
        'edit_url': 'edit_maintenance_task',
        'delete_url':'delete_maintenance_task',
        'notifications':notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'maintenance_task_list.html', context)

@user_passes_test(lambda u: is_mo(u), login_url=None)
def add_maintenance_task_page(request):
    user_branch = request.user.userprofile.branch
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()


    return render(request, 'add_maintenance_task.html',{ 'active_page': 'maintenance_task_list', 'notifications':notifications,'latest_notification_id': latest_notification.id if latest_notification else 0,
})


@user_passes_test(lambda u: is_mo(u), login_url=None)
def add_maintenance_task(request):
    user_branch = request.user.userprofile.branch
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    if request.method == 'POST':
        maintenance_task_form = MaintenanceTaskForm(request.POST)
        if maintenance_task_form.is_valid():
            maintenance_task = maintenance_task_form.save()

            frequencies = ['daily', 'weekly', 'monthly','quarterly', 'biannual', 'annual']
            for frequency in frequencies:
                task_descriptions = request.POST.getlist(f'{frequency}_tasks[]')
                if task_descriptions:
                    task_group = TaskGroup.objects.create(maintenance_task=maintenance_task, frequency=frequency)
                    for description in task_descriptions:
                        Task.objects.create(task_group=task_group, description=description)

            messages.success(request, 'Maintenance Task and associated tasks added successfully!')
            return redirect('maintenance_task_list')
    else:
        maintenance_task_form = MaintenanceTaskForm()

    return render(request, 'add_maintenance_task.html', {
        'active_page': 'maintenance_task_list',
        'branch': user_branch,
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
        'maintenance_task_form': maintenance_task_form,
        'frequencies': ['daily', 'weekly', 'monthly','quarterly', 'biannual', 'annual'],
    })
@user_passes_test(lambda u: is_mo(u) or is_md(u) or is_tec(u), login_url=None)

def edit_maintenance_task(request, id):
    user_branch = request.user.userprofile.branch
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()
    maintenance_task = get_object_or_404(MaintenanceTask, id=id)

    if request.method == 'POST':
        form = MaintenanceTaskForm(request.POST, instance=maintenance_task)
        if form.is_valid():
            maintenance_task = form.save()

            # Update tasks for each frequency
            frequencies = ['daily', 'weekly', 'monthly','quarterly', 'biannual', 'annual']
            for frequency in frequencies:
                task_descriptions = request.POST.getlist(f'{frequency}_tasks[]')
                task_group, created = TaskGroup.objects.get_or_create(
                    maintenance_task=maintenance_task,
                    frequency=frequency
                )

                # Get existing tasks for this frequency
                existing_tasks = task_group.tasks.all()

                # Update or create tasks
                new_tasks = []
                for description in task_descriptions:
                    if description.strip():  # Ignore empty tasks
                        # Check if a task with this description already exists
                        task = existing_tasks.filter(description=description).first()
                        if not task:
                            # Create a new task if it doesn't exist
                            task = Task.objects.create(task_group=task_group, description=description)
                        new_tasks.append(task)

                # Delete tasks that are no longer needed
                tasks_to_delete = existing_tasks.exclude(id__in=[task.id for task in new_tasks])
                tasks_to_delete.delete()

            messages.success(request, 'Maintenance Task updated successfully!')
            return redirect('maintenance_task_list')
    else:
        form = MaintenanceTaskForm(instance=maintenance_task)

    # Preprocess tasks for each frequency
    frequencies = ['daily', 'weekly', 'monthly','quarterly', 'biannual', 'annual']
    task_groups = {}
    for frequency in frequencies:
        task_group = maintenance_task.task_groups.filter(frequency=frequency).first()
        if task_group:
            task_groups[frequency] = task_group.tasks.all()
        else:
            task_groups[frequency] = []

    return render(request, 'edit_maintenance_task.html', {
        'form': form,
        'maintenance_task': maintenance_task,
        'frequencies': frequencies,
        'task_groups': task_groups,
        'active_page': 'maintenance_task_list',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    })

@user_passes_test(lambda u: is_mo(u), login_url=None)
def delete_maintenance_task(request, id):
    # Fetch the MaintenanceTask instance or return a 404 error if not found
    maintenancetask = get_object_or_404(MaintenanceTask, id=id)
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()
    
    if request.method == 'POST':
        # Check if any Equipment objects reference this MaintenanceTask's equipment_type
        related_equipment = Equipment.objects.filter(equipment_type=maintenancetask.equipment_type)
        
        if related_equipment.exists():
            # If there are related Equipment objects, render the protected page
            messages.error(request, 'This maintenance task cannot be deleted because there are equipments with this type of maitnenance task.')
            return render(request, 'confirm_delete_protected.html', {
                'active_page': 'maintenance_task_list',
                'object': maintenancetask,
                'protected_objects': related_equipment,  # Pass the related Equipment objects
                'model_name': 'Maintenance Task',  # Dynamic model name
                'notifications': notifications,
                'latest_notification_id': latest_notification.id if latest_notification else 0,
            })
        
        # If no related Equipment objects, delete the MaintenanceTask
        maintenancetask.delete()
        messages.success(request, 'The maintenance task was deleted successfully.')
        return redirect(reverse('maintenance_task_list'))  # Redirect to the list view
    
    # If it's a GET request, render the confirmation page
    return render(request, 'confirm_delete.html', {
        'active_page': 'maintenance_task_list',
        'object': maintenancetask,
        'model_name': 'Maintenance Task',  # Dynamic model name
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    })
#-----------------------------------------------------------------add task for frequency----------------------------------------


def add_tasks_for_frequency(request, frequency):
    user_branch = request.user.userprofile.branch
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    if request.method == 'POST':
        # Get the maintenance task ID from the form (assuming it's passed as a hidden field)
        maintenance_task_id = request.POST.get('maintenance_task_id')
        if not maintenance_task_id:
            messages.error(request, 'Maintenance Task ID is missing.')
            return redirect('add_maintenance_task')

        # Fetch the maintenance task
        try:
            maintenance_task = MaintenanceTask.objects.get(id=maintenance_task_id)
        except MaintenanceTask.DoesNotExist:
            messages.error(request, 'Invalid Maintenance Task.')
            return redirect('add_maintenance_task')

        # Create or fetch the TaskGroup for the given frequency
        task_group, created = TaskGroup.objects.get_or_create(
            maintenance_task=maintenance_task,
            frequency=frequency
        )

        # Get the list of tasks from the form
        task_descriptions = request.POST.getlist('tasks[]')
        if not task_descriptions:
            messages.error(request, 'No tasks provided.')
            return redirect('add_maintenance_task')

        # Save each task
        for description in task_descriptions:
            Task.objects.create(task_group=task_group, description=description)

        messages.success(request, f'{frequency.capitalize()} tasks added successfully!')
        return redirect('add_maintenance_task')

    else:
        # Render the form for adding tasks
        maintenance_task_form = MaintenanceTaskForm()
        task_form = TaskForm()

    return render(request, 'add_tasks_for_frequency.html', {
        'active_page': 'maintenance_task_list',
        'branch': user_branch,
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
        'maintenance_task_form': maintenance_task_form,
        'task_form': task_form,
        'frequency': frequency,

    })


@csrf_exempt
def add_task(request, frequency):
    if request.method == 'POST':
        data = json.loads(request.body)
        description = data.get('description')

        # Fetch the maintenance task (you may need to pass the ID in the request)
        maintenance_task = MaintenanceTask.objects.get(id=1)  # Replace with your logic

        # Create or fetch the TaskGroup
        task_group, created = TaskGroup.objects.get_or_create(
            maintenance_task=maintenance_task,
            frequency=frequency
        )

        # Create the task
        Task.objects.create(task_group=task_group, description=description)

        return JsonResponse({'success': True})
    return JsonResponse({'success': False})

@csrf_exempt
def delete_task(request, task_id):
    if request.method == 'POST':
        task = Task.objects.get(id=task_id)
        task.delete()
        return JsonResponse({'success': True})
    return JsonResponse({'success': False})
#--------------------------------------------------------------------------------------------

@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_mo(u), login_url=None)
def equipment_list(request):
    user_branch = request.user.userprofile.branch
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    # Fetch all branches for the branch filter (only for Maintenance Oversight)
    branches = Branch.objects.all()

    # Fetch all equipment types from the MaintenanceTask model
    equipment_types = MaintenanceTask.objects.values_list('equipment_type', flat=True).distinct()

    # Filter equipment based on user role
    if request.user.userprofile.role in ['MO', 'Maintenance Oversight']:
        equipments = Equipment.objects.filter(decommissioned = False)  # Show all equipment for MO
    else:
        equipments = Equipment.objects.filter(branch=user_branch, decommissioned = False)  # Filter by branch for other roles

    context = {
        'active_page': 'equipment_list',
        'title': 'Equipments',
        'item_list': equipments,
        'edit_url': 'edit_equipment',
        'delete_url': 'delete_equipment',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
        'branches': branches,
        'equipment_types': equipment_types,
    }
    return render(request, 'equipment_list.html', context)
@user_passes_test(lambda u: is_tec(u) or is_md(u) , login_url=None)
def add_equipment_page(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()
    user_branch = request.user.userprofile.branch

    # Fetch unique equipment types from MaintenanceTask
    equipment_types = MaintenanceTask.objects.values_list('equipment_type', flat=True).distinct()

    manufacturers = Manufacturer.objects.filter(site=user_branch)
    branch = user_branch
    return render(request, 'equipment.html', {
        'manufacturers': manufacturers,
        'branch': branch,
        'equipment_types': equipment_types,  # Pass equipment types to the template
        'active_page': 'equipment_list',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    })

@user_passes_test(lambda u: is_tec(u) or is_md(u) , login_url=None)
def add_equipment(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()
    user_branch = request.user.userprofile.branch
    manufacturers = Manufacturer.objects.filter(site=user_branch)

    # Fetch unique equipment types from MaintenanceTask
    equipment_types = MaintenanceTask.objects.values_list('equipment_type', flat=True).distinct()

    if request.method == 'POST':
        # Get form data
        name = request.POST.get('name')
        equipment_type = request.POST.get('equipment_type')
        manufacturer = request.POST.get('manufacturer')
        model_number = request.POST.get('model_number')
        serial_number = request.POST.get('serial_number')
        location = request.POST.get('location')
        installation_date = request.POST.get('installation_date')
        # maintenance_interval_years = request.POST.get('maintenance_interval_years', '0')
        # maintenance_interval_months = request.POST.get('maintenance_interval_months', '0')
        # maintenance_interval_weeks = request.POST.get('maintenance_interval_weeks', '0')
        # maintenance_interval_days = request.POST.get('maintenance_interval_days', '0')
        status = request.POST.get('status')
        remark = request.POST.get('remark')

        # Validate required fields
        if not name or not equipment_type or not model_number or not serial_number or not location or not installation_date or not status:
            messages.error(request, 'Please fill out all required fields.')
            context = {
                'manufacturers': manufacturers,
                'branch': user_branch,
                'equipment_types': equipment_types,  # Pass equipment types to the template
                'active_page': 'equipment_list',
                'notifications': notifications,
                'latest_notification_id': latest_notification.id if latest_notification else 0,
            }
            return render(request, 'equipment.html', context)

        try:
            # Convert dates from string to date objects
            installation_date = datetime.strptime(installation_date, '%Y-%m-%d').date()

            # Convert maintenance intervals to integers
            # maintenance_interval_years = int(maintenance_interval_years)
            # maintenance_interval_months = int(maintenance_interval_months)
            # maintenance_interval_weeks = int(maintenance_interval_weeks)
            # maintenance_interval_days = int(maintenance_interval_days)

            # Create and save the Equipment object
            Equipment.objects.create(
                name=name,
                equipment_type=equipment_type,
                manufacturer=manufacturer,
                model_number=model_number,
                serial_number=serial_number,
                branch=user_branch,
                location=location,
                installation_date=installation_date,
                # maintenance_interval_years=maintenance_interval_years,
                # maintenance_interval_months=maintenance_interval_months,
                # maintenance_interval_weeks=maintenance_interval_weeks,
                # maintenance_interval_days=maintenance_interval_days,
                status=status,
                remark=remark
            )
            messages.success(request, 'Equipment added successfully!')
            return redirect('equipment_list')
        except ValueError as e:
            messages.error(request, f'Invalid date format: {str(e)}')
            context = {
                'manufacturers': manufacturers,
                'branch': user_branch,
                'equipment_types': equipment_types,  # Pass equipment types to the template
                'active_page': 'equipment_list',
                'notifications': notifications,
                'latest_notification_id': latest_notification.id if latest_notification else 0,
            }
            return render(request, 'equipment.html', context)
        except Exception as e:
            messages.error(request, f'An error occurred: {str(e)}')
            context = {
                'manufacturers': manufacturers,
                'branch': user_branch,
                'equipment_types': equipment_types,  # Pass equipment types to the template
                'active_page': 'equipment_list',
                'notifications': notifications,
                'latest_notification_id': latest_notification.id if latest_notification else 0,
            }
            return render(request, 'equipment.html', context)

    # For GET requests, render the form with manufacturers and the user's branch
    context = {
        'manufacturers': manufacturers,
        'branch': user_branch,
        'equipment_types': equipment_types,  # Pass equipment types to the template
        'active_page': 'equipment_list',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'equipment.html', context)

@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_mo(u) , login_url=None)
def edit_equipment(request, id):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()
    equipment = get_object_or_404(Equipment, id=id)

    # Fetch unique equipment types from MaintenanceTask
    equipment_types = MaintenanceTask.objects.values_list('equipment_type', flat=True).distinct()
    maintenance_history= MaintenanceRecord.objects.filter(
        equipment=equipment
    ).exclude(
        Q(remark__exact='') & Q(problems__exact='')
    ).order_by('-created_at')
    
    workorder_history= WorkOrder.objects.filter(
        equipment=equipment
    ).order_by('-created_at')
    # Initialize the form variable
    form = EquipmentForm(instance=equipment)

    if request.method == 'POST':
        if request.user.userprofile.role in ["MD manager", "TEC"]:
            form = EquipmentForm(request.POST, instance=equipment)
            if form.is_valid():
                form.save()
                messages.success(request, 'Equipment changed successfully')
                return redirect('equipment_list')
        else:
            messages.error(request, 'You do not have the permission to perform this task')

    return render(request, 'edit_equipment.html', {
        'form': form,
        'equipment': equipment,
        'equipment_types': equipment_types,  # Pass equipment types to the template
        'maintenance_history':maintenance_history,
        'workorder_history':workorder_history,
        'active_page': 'equipment_list',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    })
@user_passes_test(lambda u: is_tec(u) or is_md(u) , login_url=None)
def delete_equipment(request, id):
    equipment = get_object_or_404(Equipment, id=id)
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    
    if request.method == 'POST':
        if request.user.userprofile.role in ["MD manager", "TEC"]:
            try:
                equipment.delete()
                messages.success(request, 'The equipment was deleted successfully.')
                return redirect(reverse('equipment_list'))
            except ProtectedError as e:
                messages.error(request, 'This equipment cannot be deleted because it is referenced by other objects.')
                return render(request, 'confirm_delete_protected.html', {
                    'object': equipment,
                    'protected_objects': e.protected_objects,
                    'model_name': 'Equipment',  # Dynamic model name
                    'active_page':'equipment_list',
                    'notifications':notifications,
                    'latest_notification_id': latest_notification.id if latest_notification else 0,
                })
        else:
            messages.error(request, f'You do not have the permsission to perform this task')
    
    return render(request, 'confirm_delete.html', {
        'object': equipment,
        'model_name': 'Equipment',  # Dynamic model name
        'active_page':'equipment_list',

        'notifications':notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,

    })

#-----------------------------------------------------------------------------------------------
@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_mo(u) or is_im(u) , login_url=None)
def spare_part_list(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    user_branch = request.user.userprofile.branch
    
    if request.user.userprofile.role in ['MO', 'Maintenance Oversight']:
         spare_parts = SparePart.objects.all()
    else:
        spare_parts = SparePart.objects.filter(branch = user_branch)
        

    context = {
        'active_page': 'spare_part_list',
        'title': 'Spare Parts',
        'item_list': spare_parts,
        'edit_url': 'edit_spare_part',  # Assuming you have an edit view set up
        'notifications':notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'spare_part_list.html', context)

@user_passes_test(lambda u: is_im(u), login_url=None)
def add_spare_part_page(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    user_branch = request.user.userprofile.branch

    branch = user_branch
    return render (request, 'add_spare_part.html', {
                'branch': branch,
                'active_page': 'spare_part_list',
                'notifications':notifications,
                'latest_notification_id': latest_notification.id if latest_notification else 0,

            })

@user_passes_test(lambda u: is_im(u) , login_url=None)
def add_spare_part(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    user_branch = request.user.userprofile.branch

    # Define context with default values
    context = {
        'branch': user_branch,
        'active_page': 'spare_part_list',
        'notifications':notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }

    if request.method == 'POST':
        # Get form data
        name = request.POST.get('name')
        branch_id = request.POST.get('branch')
        store = request.POST.get('store')
        quantity = request.POST.get('quantity')
        part_number = request.POST.get('part_number')
        price = request.POST.get('price')
        description = request.POST.get('description')

        # Validate required fields
        if not name or not branch_id or not store or not quantity or not part_number or not price:
            messages.error(request, 'Please fill out all required fields.')
            return render(request, 'add_spare_part.html', context)

        try:
            # Convert quantity and price to appropriate types
            quantity = int(quantity)
            price = float(price)

            # Check if the part number already exists for the selected branch
            

            # Create and save the SparePart object
            SparePart.objects.create(
                name=name,
                branch_id=branch_id,
                store=store,
                quantity=quantity,
                part_number=part_number,
                price=price,
                description=description
            )
            messages.success(request, 'Spare part added successfully!')
            return redirect('spare_part_list')  # Redirect to the spare part list page
        except ValueError as e:
            messages.error(request, f'Invalid input: {str(e)}')
            return render(request, 'add_spare_part.html', context)
        except Exception as e:
            messages.error(request, f'An error occurred: {str(e)}')
            return render(request, 'add_spare_part.html', context)

    # For GET requests, render the form with branches
    return render(request, 'add_spare_part.html', context)

@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_mo(u) or is_im(u), login_url=None)
def edit_spare_part(request, id):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()
    spare_part = get_object_or_404(SparePart, id=id)  # Fetch the spare part instance

    # Initialize the form variable
    form = SparePartForm(instance=spare_part)

    if request.method == 'POST':
        if request.user.userprofile.role in ["IM"]:
            form = SparePartForm(request.POST, instance=spare_part)
            if form.is_valid():
                form.save()  # Save the changes to the database
                messages.success(request, 'Spare part updated successfully!')
                return redirect('spare_part_list')  # Redirect to the spare part list after saving
        else:
            messages.error(request, 'You do not have the permission to perform this task')

    return render(request, 'edit_spare_part.html', {
        'form': form,
        'spare_part': spare_part,
        'active_page': 'spare_part_list',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    })

 #-----------------------------------------------maintenance-----------------------------------------------------------
    
        

    
@user_passes_test(lambda u: is_tec(u) or is_md(u), login_url=None)
def add_maintenance_page(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    user_branch = request.user.userprofile.branch

    context = {
        'equipments': Equipment.objects.filter(branch=user_branch, decommissioned = False),
        'technicians': User.objects.filter(
            userprofile__branch=user_branch,  # Filter by branch
            userprofile__role='TEC'  # Filter by role (Technician)
        ),
        'branch': user_branch,
        'maintenance_tasks': MaintenanceTask.objects.all(),
        'work_orders': WorkOrder.objects.filter(branch=user_branch),
        'spare_parts': SparePart.objects.filter(branch=user_branch),
        'active_page': 'maintenance_list',
        'notifications':notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'add_maintenance.html', context)

@user_passes_test(lambda u: is_tec(u) or is_md(u), login_url=None)
def add_maintenance(request):
    user_branch = request.user.userprofile.branch
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()
    spare_parts = SparePart.objects.filter(branch=user_branch)
    issued_spare_parts = SparePartRequest.objects.filter(
        technician=request.user,
        status='Issued',
        spare_part__branch=user_branch
    ).values_list('spare_part_id', flat=True)
    spare_parts = spare_parts.filter(id__in=issued_spare_parts)
    context = {
        'equipments': Equipment.objects.filter(branch=user_branch, decommissioned = False),
        'technicians': User.objects.filter(userprofile__branch=user_branch, userprofile__role='TEC'),
        'branch': user_branch,
        'maintenance_tasks': MaintenanceTask.objects.all(),
        'work_orders': WorkOrder.objects.filter(branch=user_branch),
        'spare_parts': spare_parts,
        'active_page': 'maintenance_list',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }

    if request.method == 'POST':
        equipment_id = request.POST.get('equipment')
        assigned_technicians = request.POST.getlist('assigned_technicians[]')
        branch_id = request.POST.get('branch')
        maintenance_type = request.POST.get('maintenance_type')  # Ensure this matches the form field name
        spare_parts = request.POST.getlist('spare_parts[]')
        spare_part_quantities = request.POST.getlist('spare_part_quantities[]')
        remark = request.POST.get('remark')
        procedure = request.POST.get('procedure')
        problems = request.POST.get('problems')
        # status = request.POST.get('status')

        try:
            # Step 1: Get the selected equipment
            equipment = Equipment.objects.get(id=equipment_id)

            # Step 2: Get the equipment_type from the selected equipment
            equipment_type = equipment.equipment_type

            # Step 3: Fetch the maintenance_task associated with the equipment_type
            maintenance_task = MaintenanceTask.objects.filter(equipment_type=equipment_type).first()

            if not maintenance_task:
                messages.error(request, f'No maintenance task found for equipment type: {equipment_type}.')
                return render(request, 'add_maintenance.html', context)

            # Step 4: Create the maintenance record
            maintenance = MaintenanceRecord.objects.create(
                equipment_id=equipment_id,
                branch_id=branch_id,
                maintenance_task=maintenance_task,
                maintenance_type=maintenance_type,  # Ensure this is correctly saved
                remark=remark,
                procedure=procedure,
                problems=problems,
                status='Not Started',
            )
            maintenance.assigned_technicians.set(assigned_technicians)

            # Step 5: Process spare parts
            for spare_part_id, quantity_used in zip(spare_parts, spare_part_quantities):
                if not spare_part_id.strip() or not quantity_used.strip():
                    continue

                spare_part = SparePart.objects.get(id=spare_part_id)
                quantity_used = int(quantity_used)

                if spare_part.quantity < quantity_used:
                    messages.error(request, f'Not enough quantity for {spare_part.name}.')
                    maintenance.delete()
                    return render(request, 'add_maintenance.html', context)

                spare_part.quantity -= quantity_used
                spare_part.save()
                check_low_spare_parts(spare_part)

                SparePartUsage.objects.create(
                    maintenance_record=maintenance,
                    spare_part=spare_part,
                    quantity_used=quantity_used,
                )

            # Step 6: Notify assigned technicians
            for technician_id in assigned_technicians:
                technician = User.objects.get(id=technician_id)
                Notification.objects.create(
                    user=technician,
                    type="maintenance",
                    message=f'You have been assigned a new maintenance task: {maintenance.equipment.name}.',
                )
                notification_created.send(sender=Notification)

            messages.success(request, 'Maintenance record added successfully!')
            return redirect('maintenance_list')
        except Exception as e:
            messages.error(request, f'An error occurred: {str(e)}')
            return render(request, 'add_maintenance.html', context)

    return render(request, 'add_maintenance.html', context)

@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_mo(u), login_url=None)
def maintenance_list(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()
    user_branch = request.user.userprofile.branch
    # Filter equipment based on user role
    if request.user.userprofile.role in ['MO', 'Maintenance Oversight']:
        maintenance_records = MaintenanceRecord.objects.all()  # Show all equipment for MO
    else:
        maintenance_records = MaintenanceRecord.objects.filter(branch = user_branch ) # Filter by branch for other roles
   
     # Fetch all maintenance records
    context = {
        'active_page': 'maintenance_list',
        'title': 'Maintenance List',
        'item_list': maintenance_records,
        'edit_url': 'edit_maintenance',  # URL name for editing maintenance records
        'delete_url':'delete_maintenance',
        'notifications':notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'maintenance_list.html', context)
    
#-----------------------------------------------------------------


@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_mo(u), login_url=None)
def edit_maintenance(request, id):
    # Get notifications and user info
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()
    user_branch = request.user.userprofile.branch
    
    # Get maintenance record and related data
    maintenance = get_object_or_404(MaintenanceRecord, id=id)
    # spare_parts = SparePart.objects.filter(branch=user_branch)
    spare_part_usages = SparePartUsage.objects.filter(maintenance_record=maintenance)
    # Get spare parts available to this technician
    spare_parts = TechnicianSparePart.objects.filter(
        technician=request.user
    ).select_related('spare_part')

    # Debug print to verify data
    print("Technician's Spare Parts:")
    for tsp in spare_parts:
        print(f"{tsp.spare_part.name} - Available: {tsp.available_quantity}")

    
    
    # Get tasks for this maintenance
    tasks = Task.objects.filter(
        task_group__maintenance_task=maintenance.maintenance_task,
        task_group__frequency=maintenance.maintenance_type
    )
    
    # Build remarks and completion status dictionary
    tasks_with_status = []
    for task in tasks:
        completion = TaskCompletion.objects.filter(
            maintenance_record=maintenance,
            task=task
        ).first()
        tasks_with_status.append({
            'task': task,
            'remark': completion.remark if completion else "",
            'is_completed': completion.is_completed if completion else False
        })

    if request.method == 'POST':
        if request.user in maintenance.assigned_technicians.all():
            # Collect all form data
            equipment_id = request.POST.get('equipment')
            assigned_technicians = request.POST.getlist('assigned_technicians')
            branch_id = request.POST.get('branch')
            maintenance_type = request.POST.get('maintenance_type')
            spare_parts_post = request.POST.getlist('spare_parts[]')
            spare_part_quantities = request.POST.getlist('spare_part_quantities[]')
            remark = request.POST.get('remark')
            procedure = request.POST.get('procedure')
            problems = request.POST.get('problems')
            status = request.POST.get('status')
            completed_tasks = request.POST.getlist('completed_tasks')

            try:
                # Step 1: Validate and get equipment
                equipment = Equipment.objects.get(id=equipment_id)
                equipment_type = equipment.equipment_type

                # Step 2: Get maintenance task for equipment type
                maintenance_task = MaintenanceTask.objects.filter(equipment_type=equipment_type).first()
                if not maintenance_task:
                    messages.error(request, f'No maintenance task found for equipment type: {equipment_type}.')
                    return render(request, 'edit_maintenance.html', {
                        'maintenance': maintenance,
                        'spare_parts': spare_parts,
                        'spare_part_usages': spare_part_usages,
                        'tasks': tasks,
                        'tasks_with_status': tasks_with_status,
                        'active_page': 'maintenance_list',
                        'notifications': notifications,
                        'latest_notification_id': latest_notification.id if latest_notification else 0,
                    })

                # Step 3: Update maintenance record fields
                maintenance.equipment_id = equipment_id
                maintenance.maintenance_task = maintenance_task
                maintenance.maintenance_type = maintenance_type
                maintenance.remark = remark
                maintenance.procedure = procedure
                maintenance.problems = problems
                maintenance.save()

                # Step 4: Handle spare parts - return old quantities first
                for usage in spare_part_usages:
                    tech_spare = TechnicianSparePart.objects.get(
                        technician=request.user,
                        spare_part=usage.spare_part
                    )
                    # spare_part = usage.spare_part
                    tech_spare.used_quantity -= usage.quantity_used
                    tech_spare.save()

                # Step 5: Process new spare part quantities
                for spare_part_id, quantity_used in zip(spare_parts_post, spare_part_quantities):
                    if not str(spare_part_id).strip() or not str(quantity_used).strip():
                        continue

                    spare_part_id = int(spare_part_id)
                    quantity_used = int(quantity_used)
                    spare_part = SparePart.objects.get(id=spare_part_id)

                    tech_spare = TechnicianSparePart.objects.get(
                        technician=request.user,
                        spare_part=spare_part
                    )
                    # request_info = spare_part_info.get(spare_part_id)
                    # Validate quantity
                    if tech_spare.available_quantity < quantity_used:
                        messages.error(request, f'Not enough quantity for {spare_part.name}. Available: {tech_spare.available_quantity}')
                        # Rollback old quantities
                        for usage in spare_part_usages:
                            # spare_part = usage.spare_part
                            tech_spare.used_quantity += usage.quantity_used
                            tech_spare.save()
                        return redirect(f'{request.path}?equipment={equipment_id}&maintenance_task={maintenance_task.id}&error=1')

                    # Update spare part quantity
                    #-------------------the double subtraction i was encountering
                    # tech_spare.used_quantity += quantity_used
                    # tech_spare.save()
                    

                    # Create/update usage record
                    SparePartUsage.objects.update_or_create(
                        maintenance_record=maintenance,
                        spare_part=spare_part,
                        defaults={'quantity_used': quantity_used},
                    )
                    req = tech_spare.request
                    SparePartTransaction.objects.update_or_create(
                       request = req, 
                       transaction_type='Usage',
                    user=request.user,
                    quantity=quantity_used,
                    notes=f'Used ({quantity_used}) in maintenance #{maintenance}'
                    )

                # Step 6: Clean up unused spare parts
                SparePartUsage.objects.filter(maintenance_record=maintenance).exclude(
                    spare_part_id__in=[int(id) for id in spare_parts_post]
                ).delete()

                # Step 7: Handle task completions and remarks
                for task in tasks:
                    task_id = str(task.id)
                    is_completed = task_id in completed_tasks
                    task_remark = request.POST.get(f'task_remarks_{task_id}', '').strip()

                    # Get or create task completion record
                    completion, created = TaskCompletion.objects.get_or_create(
                        maintenance_record=maintenance,
                        task=task,
                        defaults={
                            'remark': task_remark,
                            'is_completed': is_completed,
                            'completed_by': request.user if is_completed else None,
                            'completed_at': timezone.now() if is_completed else None
                        }
                    )

                    # Update existing record if needed
                    if not created:
                        update_fields = []
                        if completion.remark != task_remark:
                            completion.remark = task_remark
                            update_fields.append('remark')
                        if completion.is_completed != is_completed:
                            completion.is_completed = is_completed
                            completion.completed_by = request.user if is_completed else None
                            completion.completed_at = timezone.now() if is_completed else None
                            update_fields.extend(['is_completed', 'completed_by', 'completed_at'])
                        
                        if update_fields:
                            completion.save(update_fields=update_fields)

                messages.success(request, 'Maintenance record updated successfully!')
                return redirect('edit_maintenance', id=id)

            except Exception as e:
                messages.error(request, f'An error occurred: {str(e)}')
                return render(request, 'edit_maintenance.html', {
                    'maintenance': maintenance,
                    'spare_parts': spare_parts,
                    'spare_part_usages': spare_part_usages,
                    'tasks': tasks,
                    'tasks_with_status': tasks_with_status,
                    'active_page': 'maintenance_list',
                    'notifications': notifications,
                    'latest_notification_id': latest_notification.id if latest_notification else 0,
                })
        else:
            messages.error(request, 'You are not assigned to this task.')

    # GET request handling
    form = MaintenanceRecordForm(instance=maintenance)
    return render(request, 'edit_maintenance.html', {
        'form': form,
        'maintenance': maintenance,
        'spare_parts': spare_parts,
        'spare_part_usages': spare_part_usages,
        'tasks': tasks,
        'tasks_with_status': tasks_with_status,
        'active_page': 'maintenance_list',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    })
#------------------------------------------------------------delete maintenance------------------------------------------
@user_passes_test(lambda u: is_tec(u) or is_md(u) , login_url=None)
def delete_maintenance(request, id):
    # Fetch the MaintenanceRecord instance or return a 404 error if not found
    maintenance_record = get_object_or_404(MaintenanceRecord, id=id)
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    
    if request.method == 'POST':
        if request.user.userprofile.role in ["MD manager", "TEC"]:
            # If the user confirms deletion, attempt to delete the record
            try:
                maintenance_record.delete()
                messages.success(request, 'The record was deleted successfully.')
                return redirect(reverse('maintenance_list'))  # Redirect to the list view
            except ProtectedError as e:
                # If there are protected objects, render the protected page
                messages.error(request, 'This record cannot be deleted because it is referenced by other objects.')
                return render(request, 'confirm_delete_protected.html', {
                    'object': maintenance_record,
                    'protected_objects': e.protected_objects,
                    'model_name': 'Maintenance Record',  # Dynamic model name
                    'active_page':'maintenance_list',
                    'notifications': notifications,
                    'latest_notification_id': latest_notification.id if latest_notification else 0,


                })
        else:
            messages.error(request, f'You do not have the permsission to perform this task')
        
    # If it's a GET request, render the confirmation page
    return render(request, 'confirm_delete.html', {
        'object': maintenance_record,
        'model_name': 'Maintenance Record',  # Dynamic model name
        'active_page':'maintenance_list',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,

    })
#-------------------------------------------------------------------get_tasks-------------------------------------
def get_tasks(request):
    equipment_id = request.GET.get('equipment_id')
    maintenance_type = request.GET.get('maintenance_type')

    if not equipment_id or not maintenance_type:
        return JsonResponse({'tasks': []})

    try:
        # Get the equipment
        equipment = Equipment.objects.get(id=equipment_id)

        # Get the maintenance task associated with the equipment's type
        maintenance_task = MaintenanceTask.objects.filter(equipment_type=equipment.equipment_type).first()

        if not maintenance_task:
            return JsonResponse({'tasks': []})

        # Get the task group for the selected maintenance type and frequency
        task_group = TaskGroup.objects.filter(
            maintenance_task=maintenance_task,
            frequency=maintenance_type
        ).first()

        if task_group:
            # Fetch tasks associated with the task group
            tasks = task_group.tasks.all()
            tasks_data = [{'id': task.id, 'description': task.description} for task in tasks]
            return JsonResponse({'tasks': tasks_data})
        else:
            return JsonResponse({'tasks': []})
    except Exception as e:
        return JsonResponse({'tasks': []})
#-----------------------------------------------------accept maintenance----------------------------------------


def accept_maintenance(request, maintenance_id):
    # Fetch the MaintenanceRecord instance or return a 404 error if not found
    maintenance = get_object_or_404(MaintenanceRecord, id=maintenance_id)
    
    # Check if the current user is assigned to the maintenance task
    if request.user in maintenance.assigned_technicians.all():
        # Update the status to 'Accepted'
        maintenance.status = 'Accepted'
        
        # Add the current user to the accepted_by ManyToManyField
        maintenance.accepted_by.add(request.user)
        
        # Save the changes
        maintenance.save()
        
        # Display a success message
        messages.success(request, 'Maintenance task accepted.')
    else:
        # Display an error message if the user is not assigned to the task
        messages.error(request, 'You are not assigned to this task.')
    
    # Redirect to the maintenance list page
    return redirect('edit_maintenance',id = maintenance_id)

#--------------------------------------complete maintenance-----------------------------------------------------------

def complete_maintenance(request, maintenance_id):
    maintenance = MaintenanceRecord.objects.get(id=maintenance_id)
    if request.user in maintenance.assigned_technicians.all():
        maintenance.status = 'Complete'
        maintenance.save()
        messages.success(request, 'Maintenance task marked as complete.')
        
        maintenance_branch = maintenance.branch
        
        # Find the MD Manager in the same branch
        managers = User.objects.filter(userprofile__branch=maintenance_branch, userprofile__role='MD manager')
        for manager in managers:
            # Create a notification for the MD Manager
            Notification.objects.create(
                user=manager,
                type = "maintenance",
                message=f'The maintenance task for {maintenance.equipment.name} has been marked as complete.',
            )
            notification_created.send(sender=Notification)
        
    else:
        messages.error(request, 'You are not assigned to this task.')
    return redirect('maintenance_list')

#-------------------------------------approve maintenance-----------------------------------------------------

def approve_maintenance(request, maintenance_id):
    maintenance = MaintenanceRecord.objects.get(id=maintenance_id)
    maintenance_type = maintenance.maintenance_type
    if request.user.userprofile.role == 'MD manager':
        maintenance.status = 'Approved'
        maintenance.approved_by = request.user
        maintenance.save()


        # Update equipment's last and next maintenance dates
        if maintenance_type == 'daily':
            equipment = maintenance.equipment
            equipment.last_daily_maintenance_date = timezone.now().date()
            equipment.next_daily_maintenance_date = calculate_next_maintenance_date(maintenance_type)
            equipment.save()
        elif maintenance_type == 'weekly':
            equipment = maintenance.equipment
            equipment.last_weekly_maintenance_date = timezone.now().date()
            equipment.next_weekly_maintenance_date = calculate_next_maintenance_date(maintenance_type)
            equipment.save()
        elif maintenance_type == 'monthly':
            equipment = maintenance.equipment
            equipment.last_monthly_maintenance_date = timezone.now().date()
            equipment.next_monthly_maintenance_date = calculate_next_maintenance_date(maintenance_type)
            equipment.save()
        elif maintenance_type == 'quarterly':
            equipment = maintenance.equipment
            equipment.last_quarterly_maintenance_date = timezone.now().date()
            equipment.next_quarterly_maintenance_date = calculate_next_maintenance_date(maintenance_type)
            equipment.save()
        elif maintenance_type == 'annual':
            equipment = maintenance.equipment
            equipment.last_annual_maintenance_date = timezone.now().date()
            equipment.next_annual_maintenance_date = calculate_next_maintenance_date(maintenance_type)
            equipment.save()
        elif maintenance_type == 'biannual':
            equipment = maintenance.equipment
            equipment.last_biannual_maintenance_date = timezone.now().date()
            equipment.next_biannual_maintenance_date = calculate_next_maintenance_date(maintenance_type)
            equipment.save()

        messages.success(request, 'Maintenance task approved.')
    else:
        messages.error(request, 'You are not authorized to approve this task.')
    return redirect('maintenance_list')

#---------------------------------------------------------------reject maintenance-----------------------------------------

def reject_maintenance(request, maintenance_id):
    
    maintenance = get_object_or_404(MaintenanceRecord, id=maintenance_id)
    if request.user.userprofile.role == 'MD manager':
        maintenance.status = 'Rejected'
        maintenance.rejected_by = request.user
        maintenance.save()
        for technician in maintenance.assigned_technicians.all():
            Notification.objects.create(
                user= technician,
                type = 'work_order',
                message=f'The {maintenance.maintenance_type} maintenance for {maintenance.equipment.name} has been rejected by {maintenance.rejected_by}.',
            )
            notification_created.send(sender=Notification)
        messages.success(request, 'maintenance Rejected.')
    else:
        messages.error(request, 'You did not assign this work order.')
    return redirect('maintenance_list')
#------------------------------------------------------accept work order------------------------------------

def accept_work_order(request, work_order_id):
    
    work_order = get_object_or_404(WorkOrder, id=work_order_id)
    if request.user.userprofile.role == 'MD manager':
        work_order.status = 'Accepted'
        work_order.accepted_by = request.user
        work_order.save()
        messages.success(request, 'Work order accepted.')
    else:
        messages.error(request, 'You are not assigned to this work order.')
    return redirect('edit_work_order', id=work_order.id)
#------------------------------------------------------reject work order---------------------------------------------
def reject_work_order(request, work_order_id):
    
    work_order = get_object_or_404(WorkOrder, id=work_order_id)
    if request.user.userprofile.role == 'MD manager':
        work_order.status = 'Rejected'
        work_order.rejected_by = request.user
        work_order.save()
        for technician in work_order.assigned_technicians.all():
            Notification.objects.create(
                user= technician,
                type = 'work_order',
                message=f'The work order for {work_order.equipment.name} has been rejected by {work_order.rejected_by}.',
            )
            notification_created.send(sender=Notification)
        messages.success(request, 'Work order Rejected.')
    else:
        messages.error(request, 'You did not assign this work order.')
    return redirect('work_order_list')

#-------------------------------------------------------complete work order--------------------

def confirm_price(request, work_order_id):
    work_order = get_object_or_404(WorkOrder, id=work_order_id)
    work_order_branch = work_order.branch
    price = work_order.price
    work_order.status = 'Price_Confirmed'
    work_order.save()
    messages.success(request, 'Work order price confirmed.')

    managers = User.objects.filter(userprofile__branch=work_order_branch, userprofile__role='MD manager')
    client = work_order.requester
    for manager in managers:
        # Create a notification for the MD Manager
        Notification.objects.create(
            user=manager,
            type = 'work_order',
            message=f'The Price for work order {work_order} of {work_order.equipment.name} has been confirmed by {client} as {price} .',
        )
        notification_created.send(sender=Notification)
    return redirect('work_order_list')

def estimate_price(request, work_order_id):
    work_order = get_object_or_404(WorkOrder, id=work_order_id)
    
    if request.user in work_order.assigned_technicians.all():
        try:
            price = Decimal(request.POST.get('price'))
            work_order.status = 'Price_Estimated'
            work_order.price = price
            work_order.save()
            
            messages.success(request, 'Work order price sent for confirmation.')
            
            # Notification logic
            work_order_branch = work_order.branch
            managers = User.objects.filter(
                userprofile__branch=work_order_branch, 
                userprofile__role='MD manager'
            )
            client = work_order.requester
            
            for manager in managers:
                Notification.objects.create(
                    user=manager,
                    type='work_order',
                    message=f'The Price for work order {work_order} of {work_order.equipment.name} has been estimated as {price}.',
                )
                notification_created.send(sender=Notification)

            Notification.objects.create(
                user=client, 
                type='work_order',
                message=f'The Price for work order {work_order} of {work_order.equipment.name} has been estimated as {price} please confirm if you agree.',
            )
            notification_created.send(sender=Notification)
            
        except (InvalidOperation, ValueError):
            messages.error(request, 'Please enter a valid price.')
    else:
        messages.error(request, 'You are not assigned to this work order.')

    return redirect('work_order_list')




def complete_work_order(request, work_order_id):
    work_order = get_object_or_404(WorkOrder, id=work_order_id)
    
    if request.user in work_order.assigned_technicians.all():
        work_order.status = 'Complete'
        work_order.save()
        
        equipment = work_order.equipment
        equipment.last_maintenance_date = timezone.now().date()
        equipment.next_maintenance_date = calculate_next_maintenance_date(equipment)
        equipment.save()
        messages.success(request, 'Work order marked as complete.')
        
        work_order_branch = work_order.branch
        
        # Find the MD Manager in the same branch
        managers = User.objects.filter(userprofile__branch=work_order_branch, userprofile__role='MD manager')
        client = work_order.requester
        for manager in managers:
            # Create a notification for the MD Manager
            Notification.objects.create(
                user=manager,
                type = 'work_order',
                message=f'The work order for {work_order.equipment.name} has been marked as complete.',
            )
            notification_created.send(sender=Notification)
        
            
    else:
        messages.error(request, 'You are not assigned to this work order.')
    
    return redirect('work_order_list')


#-------------------------------------------------------approve work order

def approve_work_order(request, work_order_id):
    work_order = get_object_or_404(WorkOrder, id=work_order_id)
    # manager = work_order.
    
# if request.user == manager:
    work_order.status = 'Approved'
    work_order.approved_by = request.user
    work_order.save()

    # Update equipment's last and next maintenance dates if relevant
    for technician in work_order.assigned_technicians.all():
        Notification.objects.create(
            user= technician,
            type = 'work_order',
            message=f'The work order for {work_order.equipment.name} has been approved by {work_order.approved_by}.',
        )
        notification_created.send(sender=Notification)

    messages.success(request, 'Work order approved.')
# else:
#     messages.error(request, 'You are not authorized to approve this work order.')

    return redirect('work_order_list')




#-----------------------------------------calculate next maintenance date--------------------------------------------------

def calculate_next_maintenance_date(maintenance_type):
    """
    Calculate the next maintenance date based on the maintenance type.
    """
    today = timezone.now().date()

    if maintenance_type == 'daily':
        # For daily maintenance, add 1 day
        return today + timedelta(days=1)
    
    elif maintenance_type == 'weekly':
        # For weekly maintenance, add 1 week
        return today + timedelta(weeks=1)
    
    elif maintenance_type == 'monthly':
        # For monthly maintenance, add 1 month (approximated as 30 days)
        return today + timedelta(days=30)
    
    elif maintenance_type == 'quarterly':
        # For quarterly maintenance, add 3 month (approximated as 91 days)
        return today + timedelta(days=91)
    
    elif maintenance_type == 'annual':
        # For annual maintenance, add 1 year (approximated as 365 days)
        return today + timedelta(days=365)
    
    elif maintenance_type == 'biannual':
        # For biannual maintenance, add 6 months (approximated as 182 days)
        return today + timedelta(days=182)
    
    else:
        # Default fallback: return today + 1 day
        return today + timedelta(days=1)

#-------------------------------------------mark notification as read--------------------------------------------------------

def mark_notification_as_read(request, notification_id):
    if request.user.is_authenticated:
        try:
            notification = Notification.objects.get(id=notification_id, user=request.user)
            notification.is_read = True
            notification.save()
            # Redirect based on the type of notification
            if notification.type == 'maintenance':  # Assuming you have a 'type' field
                return redirect('maintenance_list')  # Adjust the URL name as necessary
            elif notification.type == 'work_order':
                return redirect('work_order_list')  # Adjust the URL name as necessary
            elif notification.type == 'low_maintenance':
                return redirect('low_maintenance')  # Adjust the URL name as necessary
            elif notification.type == 'low_spare_part':
                return redirect('low_spare_part')  # Adjust the URL name as necessary
            elif notification.type == 'expiration_list':
                return redirect('expired_chemical.html')
            elif notification.type == 'maintenance_due':
                return redirect('maintenance_due')
            elif notification.type == 'spare_part_request':
                return redirect('issue_list')
            
        except Notification.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Notification not found'})
    return JsonResponse({'status': 'error', 'message': 'User not authenticated'})










def check_low_spare_parts(spare_part):
    """
    Check if the spare part quantity is below 5 and create a notification for the MD manager.
    """
    if spare_part.quantity < 5:
        # Get the MD manager of the same branch
        md_managers = User.objects.filter(
            userprofile__branch=spare_part.branch,
            userprofile__role='MD manager'
        )

        for md_manager in md_managers:
            # Create a notification for the MD manager
            Notification.objects.create(
                user=md_manager,
                type="low_spare_part",
                message=f'Low stock alert: {spare_part.name} is below 5 in {spare_part.branch.name}.',
            )
            notification_created.send(sender=Notification)

def low_spare_part(request):
    # Get the user's branch
    user_branch = request.user.userprofile.branch
    
    if request.user.userprofile.role in ['MO', 'Maintenance Oversight']:
        spare_parts = SparePart.objects.filter(quantity__lt=5)

    else:
        spare_parts = SparePart.objects.filter(branch=user_branch, quantity__lt=5)  

    # Fetch spare parts with quantity below 5 for the user's branch

    # Get notifications for the user
    notifications = Notification.objects.filter(user=request.user, is_read=False).order_by('-timestamp')[:10]
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    context = {
        'spare_parts': spare_parts,  # Changed variable name to match template
        'active_page': 'low_spare_part',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }

    return render(request, 'low_spare_part.html', context)


@user_passes_test(lambda u: is_im(u), login_url=None)
def restock_spare_part(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()
    form = RestockSparePartForm()
    if request.method == 'POST':
        if request.user.userprofile.role in ["IM"]:
            selected_spare_part_id = request.POST.get('spare_part')  # Get the selected spare part ID
            spare_part = get_object_or_404(SparePart, id=selected_spare_part_id)  # Fetch the spare part instance

            form = RestockSparePartForm(request.POST, request.FILES)
            if form.is_valid():
                restock = form.save(commit=False)
                restock.spare_part = spare_part  # Associate the restock with the spare part
                restock.restock_date = timezone.now()  # Set the restock date to the current time
                restock.save()

                # Update the spare part quantity and last_restock_date
                spare_part.quantity += restock.quantity
                spare_part.last_restock_date = restock.restock_date
                spare_part.save()

                messages.success(request, f'{restock.quantity} units of {spare_part.name} restocked successfully!')
                return redirect('spare_part_list')
        else:
            messages.error(request, f'You do not have the permsission to perform this task')
    else:
        form = RestockSparePartForm()

    # Pass a list of spare parts to the template
    user_branch = request.user.userprofile.branch
    spare_parts = SparePart.objects.filter(branch=user_branch)

    context = {
        'form': form,
        'spare_parts': spare_parts,  # Ensure you pass all spare parts to the template
        'active_page': 'restock_list',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'restock_spare_part.html', context)

@user_passes_test(lambda u: is_tec(u) or is_md(u) or is_mo(u) or is_im(u), login_url=None)
def restock_list(request):
    # Get the user's branch
    user_branch = request.user.userprofile.branch
    
    if request.user.userprofile.role in ['MO', 'Maintenance Oversight']:
        restock_list = RestockSparePart.objects.all().order_by('-restock_date')

    else:
        restock_list = RestockSparePart.objects.filter(spare_part__branch=user_branch).order_by('-restock_date')

    # Fetch restock records for spare parts in the user's branch

    # Get notifications for the user
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    context = {
        'restock_list': restock_list,
        'active_page': 'restock_list',
        'notifications': notifications,
        'title': 'Restock List',
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }

    return render(request, 'restock_list.html', context)

@user_passes_test(lambda u: is_im(u), login_url=None)
def restock_spare_part_page(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    user_branch = request.user.userprofile.branch

    context = {
        'spare_parts': SparePart.objects.filter(branch=user_branch),  # Fetch spare parts for the branch
        'active_page': 'restock_list',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }

    return render(request, 'restock_spare_part.html', context)


def maintenance_due(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    today = timezone.now().date()
    due_in_5_days = today + timezone.timedelta(days=5)
    print(f"Today: {today}, Due in 5 days: {due_in_5_days}")

    # Fetch equipment due for maintenance within the next 5 days
    due_equipment = Equipment.objects.filter(
        next_monthly_maintenance_date__lte=due_in_5_days
    ) | Equipment.objects.filter(
        next_quarterly_maintenance_date__lte=due_in_5_days
    ) | Equipment.objects.filter(
        next_biannual_maintenance_date__lte=due_in_5_days
    ) | Equipment.objects.filter(
        next_annual_maintenance_date__lte=due_in_5_days
    )
    # print(f"Due equipment before exclusion: {due_equipment}")

    # Role-based filtering
    user_role = request.user.userprofile.role  # Assuming the user's role is stored in a `role` field on the User model
    if user_role in ['MD manager', 'TEC']:  # Manager or Technician
        due_equipment = due_equipment.filter(branch=request.user.userprofile.branch)  # Filter by the user's branch
    elif user_role == 'MO':  # Maintenance Oversight
        # No additional filtering needed; show all due equipment
        pass
    else:
        # For other roles, return an empty list or restrict access
        due_equipment = Equipment.objects.none()

    context = {
        'due_equipment': due_equipment,
        'due_in_5_days': due_in_5_days,  # Pass the due_in_5_days to the template
        'active_page': 'maintenance_due',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'maintenance_due.html', context)
#----------------------------------------------------Import Export-------------------------------------------
def export_data(request, model_name):
    """
    Generalized view for exporting data in Excel format (XLSX) with branch filtering.
    """
    # Map model names to their respective resources
    model_resource_map = {
        'Equipment': EquipmentResource,
    }

    if model_name not in model_resource_map:
        return HttpResponse("Model not supported.", status=400)

    # Get the resource class
    resource_class = model_resource_map[model_name]
    resource = resource_class()

    # Get the user's profile
    user_profile = request.user.userprofile
    is_mo_user = user_profile.role == 'MO'  # Maintenance Oversight user

    # Apply branch filtering if user is not MO
    if not is_mo_user:
        if model_name == 'Equipment':
            queryset = Equipment.objects.filter(branch=user_profile.branch)
        elif model_name == 'MaintenanceRecord':
            queryset = MaintenanceRecord.objects.filter(equipment__branch=user_profile.branch)
        dataset = resource.export(queryset)
    else:
        # For MO users, export all data
        dataset = resource.export()

    # Export the data in Excel format (XLSX)
    file_format = base_formats.XLSX()
    response = HttpResponse(
        file_format.export_data(dataset),
        content_type=file_format.get_content_type()
    )
    response['Content-Disposition'] = f'attachment; filename={model_name}_export.xlsx'
    return response

#----------------------------------------------------------import--------------------------------------------------

def import_data(request, model_name):
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    """
    Generalized view for importing data with a preview page.
    """
    # Map model names to their respective resources
    notifications = get_notifications(request.user)
    model_resource_map = {
        'Equipment': EquipmentResource,  # Add more models here
        # Example: 'Customer': CustomerResource,
    }

    if model_name not in model_resource_map:
        messages.error(request, "Model not supported.")
        return redirect('home')  # Redirect to a safe page
    
    try:
        user_profile = request.user.userprofile
        user_branch = user_profile.branch
        is_mo_user = user_profile.role == 'MO'  # Maintenance Oversight
    except AttributeError:
        messages.error(request, "User profile not configured properly.")
        return redirect('home')

    # Block MO users from importing - ADDED THIS BLOCK
    if is_mo_user:
        messages.error(request, "You don't have permission to import data.")
        return redirect('equipment_list')

    # Get the resource class
    resource_class = model_resource_map[model_name]
    resource = resource_class()

    if request.method == 'POST':
        if 'confirm_import' in request.POST:
            # Retrieve the temporary file path from the session
            temp_file_path = request.session.get('temp_file_path')
            if not temp_file_path or not os.path.exists(temp_file_path):
                messages.error(request, "File not found. Please upload the file again.")
                return redirect('import_data', model_name=model_name)

            try:
                # Read the file content from the temporary file
                with open(temp_file_path, 'rb') as f:
                    file_content = f.read()

                # Load the dataset
                file_format = base_formats.XLSX()
                dataset = file_format.create_dataset(file_content)
                result = resource.import_data(dataset, dry_run=False, user_branch=user_branch)  # Perform the actual import

                if not result.has_errors():
                    messages.success(request, 'Data imported successfully.')
                else:
                    # Collect errors and display them
                    error_messages = []
                    for row in result.invalid_rows:
                        error_messages.append(f"Row {row.number}: {row.error}")
                    for row in result.row_errors():
                        for error in row[1]:
                            error_messages.append(f"Row {row[0]}: {error.error}")

                    # Pass errors to the template
                    return render(request, 'import_template.html', {
                        'model_name': model_name,
                        'errors': error_messages,
                        'active_page': 'equipment_list',
                        'notifications': notifications,
                    })

                # Clean up the temporary file
                os.remove(temp_file_path)
                del request.session['temp_file_path']

                return redirect('equipment_list')  # Redirect to the equipment list page

            except Exception as e:
                messages.error(request, f"An error occurred during import: {str(e)}")
                return redirect('import_data', model_name=model_name)

        elif 'import_file' in request.FILES:
            # Perform a dry-run to preview the data
            import_file = request.FILES['import_file']
            file_format = base_formats.XLSX()  # Only accept Excel files

            # Check if the file is an Excel file
            if not import_file.name.endswith('.xlsx'):
                messages.error(request, "Please upload an Excel file (.xlsx).")
                return redirect('import_data', model_name=model_name)

            try:
                # Save the file temporarily
                with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as temp_file:
                    for chunk in import_file.chunks():
                        temp_file.write(chunk)
                    temp_file_path = temp_file.name

                # Store the temporary file path in the session
                request.session['temp_file_path'] = temp_file_path

                # Read the file content as bytes
                with open(temp_file_path, 'rb') as f:
                    file_content = f.read()

                # Load the dataset
                dataset = file_format.create_dataset(file_content)
                result = resource.import_data(dataset, dry_run=True, user_branch=user_branch, raise_errors=False )  # Perform a dry-run

                # Prepare data for the preview
                preview_data = []
                for row in dataset.dict:
                    preview_data.append(row)

                # Collect errors (if any)
                error_messages = []
                for row in result.invalid_rows:
                    error_messages.append(f"Row {row.number}: {row.error}")
                for row in result.row_errors():
                    for error in row[1]:
                        error_messages.append(f"Row {row[0]}: {error.error}")

                # If there are errors, do not show the confirm button
                show_confirm_button = not error_messages

                return render(request, 'import_preview.html', {
                    'model_name': model_name,
                    'preview_data': preview_data,
                    'errors': error_messages,
                    'show_confirm_button': show_confirm_button,
                    'active_page': 'equipment_list',
                    'notifications': notifications,
                    'latest_notification_id': latest_notification.id if latest_notification else 0,
                })

            except Exception as e:
                messages.error(request, f"An error occurred during preview: {str(e)}")
                return redirect('import_data', model_name=model_name)

    return render(request, 'import_template.html', {
        'model_name': model_name,
        'active_page': 'equipment_list',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    })
#------------------------------------------------------------------------------------------------------------
@user_passes_test(lambda u: is_tec(u) or is_md(u), login_url=None)
def maintenance_dashboard(request):
    user = request.user
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    user_role = user.userprofile.role
    user_branch = user.userprofile.branch
    equipment_list = Equipment.objects.filter(branch=user_branch).order_by('name')
    
    # Get date range from request
    from_date = request.GET.get('from_date', (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d'))
    to_date = request.GET.get('to_date', datetime.now().strftime('%Y-%m-%d'))

    # Equipment status counts
    operational_count = Equipment.objects.filter(
        branch=user_branch,
        status='operational',
        created_at__range=[from_date, to_date]
    ).count()
    non_operational_count = Equipment.objects.filter(
        branch=user_branch,
        status='non_operational',
        created_at__range=[from_date, to_date]
    ).count()
    under_maintenance_count = Equipment.objects.filter(
        branch=user_branch,
        status='under_maintenance',
        created_at__range=[from_date, to_date]
    ).count()

    # Maintenance by frequency counts
    daily_maintenance_count = MaintenanceRecord.objects.filter(
        branch=user_branch,
        maintenance_type='daily',
        datetime__range=[from_date, to_date]
    ).count()
    weekly_maintenance_count = MaintenanceRecord.objects.filter(
        branch=user_branch,
        maintenance_type='weekly',
        datetime__range=[from_date, to_date]
    ).count()
    monthly_maintenance_count = MaintenanceRecord.objects.filter(
        branch=user_branch,
        maintenance_type='monthly',
        datetime__range=[from_date, to_date]
    ).count()
    quarterly_maintenance_count = MaintenanceRecord.objects.filter(
        branch=user_branch,
        maintenance_type='quarterly',
        datetime__range=[from_date, to_date]
    ).count()
    biannual_maintenance_count = MaintenanceRecord.objects.filter(
        branch=user_branch,
        maintenance_type='biannual',
        datetime__range=[from_date, to_date]
    ).count()
    annual_maintenance_count = MaintenanceRecord.objects.filter(
        branch=user_branch,
        maintenance_type='annual',
        datetime__range=[from_date, to_date]
    ).count()

    # Work orders counts
    pending_work_orders = WorkOrder.objects.filter(
        branch=user_branch,
        status='Pending',
        created_at__range=[from_date, to_date]
    ).count()
    completed_work_orders = WorkOrder.objects.filter(
        branch=user_branch,
        status='Approved',
        created_at__range=[from_date, to_date]
    ).count()

    # Maintenance by month data
    maintenance_months = []
    maintenance_by_month_data = []
    for i in range(1, 13):  # January to December
        month = datetime(datetime.now().year, i, 1).strftime('%b')
        maintenance_months.append(month)
        maintenance_count = MaintenanceRecord.objects.filter(
            branch=user_branch,
            datetime__month=i,
            datetime__year=datetime.now().year
        ).count()
        maintenance_by_month_data.append(maintenance_count)

    # Work orders by month data
    work_order_months = []
    pending_work_orders_by_month = []
    completed_work_orders_by_month = []
    for i in range(1, 13):  # January to December
        month = datetime(datetime.now().year, i, 1).strftime('%b')
        work_order_months.append(month)
        pending_count = WorkOrder.objects.filter(
            branch=user_branch,
            created_at__month=i,
            created_at__year=datetime.now().year,
            status='Pending'
        ).count()
        completed_count = WorkOrder.objects.filter(
            branch=user_branch,
            created_at__month=i,
            created_at__year=datetime.now().year,
            status='Approved'
        ).count()
        pending_work_orders_by_month.append(pending_count)
        completed_work_orders_by_month.append(completed_count)

    # Spare part usage data
    spare_part_usage = SparePartUsage.objects.filter(
        (Q(maintenance_record__branch=user_branch) | Q(work_order__branch=user_branch)),
        created_at__range=[from_date, to_date]
    ).values('spare_part__name').annotate(total_used=Sum('quantity_used'))

    # Extract labels and data
    spare_part_labels = [item['spare_part__name'] for item in spare_part_usage if item['spare_part__name']]
    spare_part_usage_data = [item['total_used'] for item in spare_part_usage if item['spare_part__name']]

    # Debugging
    print("Spare Part Labels:", spare_part_labels)
    print("Spare Part Usage Data:", spare_part_usage_data)

    # Fallback for empty data
    if not spare_part_labels:
        spare_part_labels = ["No Data"]
        spare_part_usage_data = [0]
    context = {
        'from_date': from_date,
        'to_date': to_date,
        'operational_count': operational_count,
        'non_operational_count': non_operational_count,
        'under_maintenance_count': under_maintenance_count,
        'daily_maintenance_count': daily_maintenance_count,
        'weekly_maintenance_count': weekly_maintenance_count,
        'monthly_maintenance_count': monthly_maintenance_count,
        'quarterly_maintenance_count': quarterly_maintenance_count,
        'biannual_maintenance_count': biannual_maintenance_count,
        'annual_maintenance_count': annual_maintenance_count,
        'pending_work_orders': pending_work_orders,
        'completed_work_orders': completed_work_orders,
        'maintenance_months': maintenance_months,
        'maintenance_by_month_data': maintenance_by_month_data,
        'work_order_months': work_order_months,
        'pending_work_orders_by_month': pending_work_orders_by_month,
        'completed_work_orders_by_month': completed_work_orders_by_month,
        'spare_part_labels': spare_part_labels,
        'spare_part_usage_data': spare_part_usage_data,
        'equipment_list' : equipment_list,
        'active_page':'dashboard',
        'notifications':notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }

    return render(request, 'maintenance_dashboard.html', context)

#----------------------------------------------------------------------------------------

@login_required
@user_passes_test(lambda u: is_im(u), login_url=None)
def inventory_dashboard(request):
    # Spare part statistics
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()
    total_spare_parts = SparePart.objects.count()
    low_stock_parts = SparePart.objects.filter(
        quantity__lte=F('min_quantity')
    ).count()
    out_of_stock_parts = SparePart.objects.filter(
        quantity=0
    ).count()
    
    # Request statistics
    pending_requests = SparePartRequest.objects.filter(
        status='Requested'
    ).count()
    approved_requests = SparePartRequest.objects.filter(
        status='Approved'
    ).count()
    issued_requests = SparePartRequest.objects.filter(
        status='Issued'
    ).count()
    
    # Recent activity
    recent_requests = SparePartRequest.objects.select_related(
        'spare_part', 'technician'
    ).order_by('-request_date')[:10]
    
    recent_transactions = SparePartTransaction.objects.select_related(
        'request', 'request__spare_part', 'user'
    ).order_by('-timestamp')[:10]
    
    # Restock alerts
    restock_alerts = SparePart.objects.annotate(
        available_quantity=F('quantity') - Coalesce(
            Sum('sparepartrequest__quantity_requested', 
                filter=Q(sparepartrequest__status__in=['Requested', 'Approved', 'Issued'])),
            0
        )
    ).filter(
        available_quantity__lte=F('min_quantity')
    ).order_by('available_quantity')
    
    # Technician inventory overview
    technician_inventory = TechnicianSparePart.objects.select_related(
        'technician', 'spare_part'
    ).values(
        'technician__username',
        'spare_part__name'
    ).annotate(
        total_received=Sum('received_quantity'),
        total_used=Sum('used_quantity'),
        remaining=Sum('received_quantity') - Sum('used_quantity')
    ).order_by('technician__username')[:10]
    
    # Monthly request analytics (last 6 months)
    six_months_ago = datetime.now() - timedelta(days=180)
    
    # Monthly request count
    monthly_requests = SparePartRequest.objects.filter(
        request_date__gte=six_months_ago
    ).annotate(
        month=TruncMonth('request_date')
    ).values(
        'month'
    ).annotate(
        count=Count('id')
    ).order_by('month')
    
    # Most requested parts by month
    most_requested_by_month = []
    for i in range(6):
        month_start = datetime.now().replace(day=1) - timedelta(days=30*i)
        month_end = month_start.replace(day=1) + timedelta(days=32)
        month_end = month_end.replace(day=1) - timedelta(days=1)
        
        top_parts = SparePartRequest.objects.filter(
            request_date__gte=month_start,
            request_date__lte=month_end
        ).values(
            'spare_part__name'
        ).annotate(
            total_requested=Sum('quantity_requested')
        ).order_by('-total_requested')[:3]
        
        if top_parts:
            most_requested_by_month.append({
                'month': month_start.strftime('%B %Y'),
                'parts': list(top_parts)
            })
    
    # Status distribution
    status_distribution = SparePartRequest.objects.values(
        'status'
    ).annotate(
        count=Count('id')
    ).order_by('-count')
    
    context = {
        'active_page': 'dashboard',
        'total_spare_parts': total_spare_parts,
        'low_stock_parts': low_stock_parts,
        'out_of_stock_parts': out_of_stock_parts,
        'pending_requests': pending_requests,
        'approved_requests': approved_requests,
        'issued_requests': issued_requests,
        'recent_requests': recent_requests,
        'recent_transactions': recent_transactions,
        'restock_alerts': restock_alerts,
        'technician_inventory': technician_inventory,
        'monthly_requests': list(monthly_requests),
        'most_requested_by_month': most_requested_by_month,
        'status_distribution': list(status_distribution),
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    
    return render(request, 'inventory_dashboard.html', context)
#----------------------------------------------------------------------------------------
# In your views.py
def equipment_maintenance_types_api(request):
    try:
        equipment_id = request.GET.get('equipment_id')
        from_date = request.GET.get('from_date')
        to_date = request.GET.get('to_date')
        
        if not equipment_id:
            return JsonResponse({'error': 'Equipment ID required'}, status=400)
            
        # Get all maintenance records for this equipment
        records = MaintenanceRecord.objects.filter(
            equipment_id=equipment_id,
            datetime__date__range=[from_date, to_date], 
            status = 'Approved'
        )
        
        # Group by month and maintenance type
        monthly_data = records.annotate(
            month=TruncMonth('datetime')
        ).values('month', 'maintenance_type').annotate(
            count=Count('id')
        ).order_by('month')
        
        # CHANGES START HERE - Properly sort months chronologically
        # First get all unique months as date objects
        month_dates = sorted(list(set(item['month'] for item in monthly_data)))
        # Then format them as strings
        months = [date.strftime('%b %Y') for date in month_dates]
        
        types = ['daily', 'weekly', 'monthly','quarterly', 'biannual', 'annual']
        
        # Prepare dataset for each maintenance type
        datasets = []
        for maint_type in types:
            type_data = []
            for month_date in month_dates:
                month_str = month_date.strftime('%b %Y')
                count = next((item['count'] for item in monthly_data 
                            if item['month'] == month_date 
                            and item['maintenance_type'] == maint_type), 0)
                type_data.append(count)
            
            datasets.append({
                'label': maint_type.capitalize(),
                'data': type_data,
                'backgroundColor': get_color_for_type(maint_type)
            })
        
        return JsonResponse({
            'labels': months,  # Now properly ordered
            'datasets': datasets
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
    

#---------------------------------------------------------------------

def equipment_maintenance_types_api_MO(request):
    try:
        equipment_id = request.GET.get('equipment_id')
        from_date = request.GET.get('from_date')
        to_date = request.GET.get('to_date')
        
        if not equipment_id:
            return JsonResponse({'error': 'Equipment ID required'}, status=400)
            
        # Get all maintenance records for this equipment
        records = MaintenanceRecord.objects.filter(
            equipment_id=equipment_id,
            datetime__date__range=[from_date, to_date], 
            status = 'Approved'
        )
        
        # Group by month and maintenance type
        monthly_data = records.annotate(
            month=TruncMonth('datetime')
        ).values('month', 'maintenance_type').annotate(
            count=Count('id')
        ).order_by('month')
        
        # CHANGES START HERE - Properly sort months chronologically
        # First get all unique months as date objects
        month_dates = sorted(list(set(item['month'] for item in monthly_data)))
        # Then format them as strings
        months = [date.strftime('%b %Y') for date in month_dates]
        
        types = ['daily', 'weekly', 'monthly','quarterly', 'biannual', 'annual']
        
        # Prepare dataset for each maintenance type
        datasets = []
        for maint_type in types:
            type_data = []
            for month_date in month_dates:
                month_str = month_date.strftime('%b %Y')
                count = next((item['count'] for item in monthly_data 
                            if item['month'] == month_date 
                            and item['maintenance_type'] == maint_type), 0)
                type_data.append(count)
            
            datasets.append({
                'label': maint_type.capitalize(),
                'data': type_data,
                'backgroundColor': get_color_for_type(maint_type)
            })
        
        return JsonResponse({
            'labels': months,  # Now properly ordered
            'datasets': datasets
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def get_color_for_type(maint_type):
    colors = {
        'daily': 'rgba(255, 99, 132, 0.7)',
        'weekly': 'rgba(54, 162, 235, 0.7)',
        'monthly': 'rgba(75, 192, 192, 0.7)',
        'quarterly':'rgba(50, 205, 50, 0.7)',
        'biannual': 'rgba(153, 102, 255, 0.7)',
        'annual': 'rgba(255, 159, 64, 0.7)'
    }
    return colors.get(maint_type, 'rgba(201, 203, 207, 0.7)')
#-----------------------------------------------------------------------------------------

def generate_report(request):
    user = request.user
    branch = request.GET.get('branch')
    from_date = request.GET.get('from_date', (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d'))
    to_date = request.GET.get('to_date', datetime.now().strftime('%Y-%m-%d'))
    # Convert the string dates to timezone-aware datetime objects
    
    report_type = request.GET.get('report_type', 'daily')
    export_format = request.GET.get('format')  # Default to PDF
    if(user.userprofile.role in ["MD manager", "TEC"]):
        user_branch = user.userprofile.branch
    elif(user.userprofile.role in ["MO"]):
        user_branch = branch
    
    if(branch != "all"):
        # Get date range and report type from request
        

        if report_type == 'work_order':
            # Fetch work orders with status 'Approved' based on date range
            work_orders = WorkOrder.objects.filter(
                branch=user_branch,
                status='Approved',
                created_at__date__range=[from_date, to_date]
            ).order_by('created_at')

            if export_format == 'docx':
                return generate_editable_doc_work_order(work_orders, report_type, from_date, to_date)
            else:
                return generate_pdf_work_order(work_orders, report_type, from_date, to_date)
        else:
            # Fetch maintenance records based on report type and date range
            maintenance_records = MaintenanceRecord.objects.filter(
                status = 'Approved',
                maintenance_type=report_type,
                branch=user_branch,
                datetime__date__range=[from_date, to_date],
                
            ).order_by('datetime')
# approved_records = MaintenanceRecord.objects.filter(status__exact='Approved', maintenance_type = 'weekly', branch = branch )
            if export_format == 'docx':
                return generate_editable_doc(maintenance_records, report_type, from_date, to_date)
            else:
                return generate_pdf(maintenance_records, report_type, from_date, to_date)
    else:
        if report_type == 'work_order':
            return generate_MO_work_order_report(from_date, to_date, export_format)
        else:
            return generate_MO_maintenance_report(from_date, to_date, export_format,report_type)
#-------------------------------------------------------------------------------------------



def generate_pdf(maintenance_records, report_type, from_date, to_date):
    # Create a PDF document
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{report_type}_maintenance_report.pdf"'

    doc = SimpleDocTemplate(response, pagesize=letter)
    elements = []
    

    # Add title
    styles = getSampleStyleSheet()
    title = Paragraph(f"{report_type.capitalize()} Maintenance Report", styles['Title'])
    elements.append(title)
    
    # Check if maintenance_records is empty
    if not maintenance_records:
        elements.append(Paragraph("No maintenance records found for the selected date range and report type.", styles['Normal']))
        doc.build(elements)
        return response
    doc.title = f"{report_type} Maintenance Report {maintenance_records[0].branch.name}"
    # Add branch and date range
    elements.append(Paragraph(f"Branch: {maintenance_records[0].branch.name}", styles['Normal']))
    elements.append(Paragraph(f"From  {from_date}  to  {to_date}", styles['Normal']))

    space_style = ParagraphStyle(name='Space', spaceAfter=30)  # Adjust spaceAfter value as needed
    elements.append(Paragraph("", space_style))  # Empty paragraph with spacing

    for i, record in enumerate(maintenance_records):
        if i > 0:
            elements.append(PageBreak())
        # Add header for each maintenance record
        elements.append(Paragraph(f"Maintenance Date: {record.datetime.strftime('%Y-%m-%d')}, Time: {record.datetime.strftime('%H:%M')}", styles['Heading5']))
        elements.append(Paragraph(f"Equipment: {record.equipment.name} (Serial: {record.equipment.serial_number}), Location: {record.equipment.location}", styles['Normal']))
        space_style = ParagraphStyle(name='Space', spaceAfter=8)  # Adjust spaceAfter value as needed
        elements.append(Paragraph("", space_style))  # Empty paragraph with spacing

        # Fetch all tasks associated with the maintenance task and maintenance type
        tasks = Task.objects.filter(
            task_group__maintenance_task=record.maintenance_task,
            task_group__frequency=record.maintenance_type
        )

        # Create table data
        data = [['No.', 'Inspection', 'Yes', 'No', 'Remarks']]
        for i, task in enumerate(tasks, start=1):
            # Check if the task was completed
            task_completion = TaskCompletion.objects.filter(maintenance_record=record, task=task).first()
            yes_mark = '✓' if task_completion and task_completion.is_completed else ''
            no_mark = '✓' if task_completion and not task_completion.is_completed else ''
            data.append([
                str(i),
                Paragraph(task.description, styles['Normal']),  # Wrap long text in a paragraph
                yes_mark,
                no_mark,
                Paragraph(task_completion.remark if task_completion else '', styles['Normal'])
            ])

        # Create table with adjusted column widths
        table = Table(data, colWidths=[0.5*inch, 4.5*inch, 0.5*inch, 0.5*inch, 2*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Align content to the top
            ('WORDWRAP', (1, 0), (1, -1), True),  # Enable word wrap for the "Inspection" column
            ('WORDWRAP', (4, 0), (4, -1), True),  # Add this line for the "Remarks" column
            ('ALIGN', (4, 0), (4, -1), 'LEFT'),  # Optional: Left align remarks text
        ]))

        elements.append(table)
        space_style = ParagraphStyle(name='Space', spaceAfter=8)  # Adjust spaceAfter value as needed
        elements.append(Paragraph("", space_style))  # Empty paragraph with spacing

        # Add inspected by and approved by sections
        inspected_by = [tech.get_full_name() for tech in record.assigned_technicians.all()]
        approved_by = record.approved_by.get_full_name() if record.approved_by else ""

        # Create a table with one row and two columns
        table_data = [
            [
                Paragraph(f"<b>Inspected by:</b><br/><br/>" + "<br/><br/>".join(inspected_by)),  # Left column
                Paragraph(f"<b>Approved by:</b><br/><br/>{approved_by}")  # Right column
            ]
        ]

        # Define column widths (adjust as needed)
        col_widths = [3 * inch, 3 * inch]  # Equal width for both columns

        # Create the table
        table = Table(table_data, colWidths=col_widths)
        table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),  # Align text to the left
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Align text to the top
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),  # Add padding at the bottom
        ]))

        # Add the table to the elements
        elements.append(table)
        # Add space between records
        elements.append(Paragraph("<br/><br/>", styles['Normal']))

    # Build the PDF
    doc.build(elements)
    return response
#-----------------------------------------------------------------------------------------

def generate_editable_doc(maintenance_records, report_type, from_date, to_date):
    # Create a Word document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{report_type}_maintenance_report.docx"'

    doc = Document()

    # Add main title
    title = doc.add_heading(f"{report_type.capitalize()} Maintenance Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    title.style.font.size = Pt(18)  # Make the title smaller
    
    if not maintenance_records:
        doc.add_paragraph("No maintenance records found for the selected date range and report type.", style='Intense Quote')
        # Save the document to a BytesIO stream
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response.write(buffer.getvalue())
        buffer.close()
        return response

    # Add branch and date range
    branch_paragraph = doc.add_paragraph(f"Branch: {maintenance_records[0].branch.name}", style='Intense Quote')
    branch_paragraph.paragraph_format.space_after = Pt(0)
    date_paragraph = doc.add_paragraph(f"From {from_date} to {to_date}", style='Intense Quote')
    date_paragraph.paragraph_format.space_before = Pt(0)  # Remove space before this paragraph
    date_paragraph.paragraph_format.space_after = Pt(20)  # Add space after this paragraph

    for record in maintenance_records:
        # Add header for each maintenance record
        doc.add_heading(f"Maintenance Date: {record.datetime.strftime('%Y-%m-%d')}, Time: {record.datetime.strftime('%H:%M')}", level=3)
        doc.add_paragraph(f"Equipment: {record.equipment.name} (Serial: {record.equipment.serial_number}), Location: {record.equipment.location}")

        # Create table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False

        # Set column widths
        col_widths = [0.2, 4, 0.2, 0.2, 2]  # Widths in inches
        for i, width in enumerate(col_widths):
            col = table.columns[i]
            col.width = Inches(width)

        # Add table headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No.'
        hdr_cells[1].text = 'Inspection'
        hdr_cells[2].text = 'Yes'
        hdr_cells[3].text = 'No'
        hdr_cells[4].text = 'Remarks'

        # Fetch all tasks associated with the maintenance task and maintenance type
        tasks = Task.objects.filter(
            task_group__maintenance_task=record.maintenance_task,
            task_group__frequency=record.maintenance_type
        )

        # Add tasks to the table
        for i, task in enumerate(tasks, start=1):
            # Check if the task was completed
            task_completion = TaskCompletion.objects.filter(maintenance_record=record, task=task).first()
            yes_mark = '✓' if task_completion and task_completion.is_completed else ''
            no_mark = '✓' if task_completion and not task_completion.is_completed else ''
            
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = task.description
            row_cells[2].text = yes_mark
            row_cells[3].text = no_mark
            remarks = task_completion.remark if task_completion and task_completion.remark else ''
            row_cells[4].text = str(remarks)

        space_paragraph = doc.add_paragraph()
        space_paragraph.paragraph_format.space_after = Pt(0.1)  # Adjust the space as needed

        # Add inspected by and approved by sections
        inspected_by = [tech.get_full_name() for tech in record.assigned_technicians.all()]
        approved_by = record.approved_by.get_full_name() if record.approved_by else ""

        # Create a table with one row and two columns
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False

        # Set column widths (adjust as needed)
        table.columns[0].width = Inches(3)  # Width for the "Inspected by" column
        table.columns[1].width = Inches(3)  # Width for the "Approved by" column

        # Add "Inspected by" to the first cell
        inspected_by_cell = table.rows[0].cells[0]
        inspected_by_cell.text = "Inspected by:"
        for paragraph in inspected_by_cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True  # Make "Inspected by" bold

        # Add technician names below "Inspected by"
        for tech in inspected_by:
            inspected_by_cell.add_paragraph(tech, style='List Bullet')

        # Add "Approved by" to the second cell
        approved_by_cell = table.rows[0].cells[1]
        approved_by_paragraph = approved_by_cell.add_paragraph()
        approved_by_paragraph.add_run("Approved by: ").bold = True  # Make "Approved by" bold
        approved_by_paragraph.add_run(approved_by)

        # Add space between records
        doc.add_paragraph()

    # Save the document to a BytesIO stream
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response.write(buffer.getvalue())
    buffer.close()

    return response
#----------------------------------------------------------------------------------------------------

def generate_pdf_work_order(work_orders, report_type, from_date, to_date):
    # Create a PDF document
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{report_type}_work_order_report.pdf"'

    doc = SimpleDocTemplate(response, pagesize=letter)
    elements = []
    

    # Add title
    styles = getSampleStyleSheet()
    title = Paragraph(f"Work Order Report", styles['Title'])
    elements.append(title)
    
    if not work_orders:
        elements.append(Paragraph("No work order records found for the selected date range.", styles['Normal']))
        doc.build(elements)
        return response
    doc.title = f"Work Order Report {work_orders[0].branch.name}"
    # Add branch and date range
    elements.append(Paragraph(f"Branch: {work_orders[0].branch.name}", styles['Normal']))
    elements.append(Paragraph(f"From  {from_date}  to  {to_date}", styles['Normal']))

    space_style = ParagraphStyle(name='Space', spaceAfter=30)  # Adjust spaceAfter value as needed
    elements.append(Paragraph("", space_style))  # Empty paragraph with spacing

    # Iterate through each work order
    for work_order in work_orders:
        # Create table data for the current work order
        elements.append(Paragraph(f"Work Order Date: {work_order.created_at.strftime('%Y-%m-%d')}", styles['Heading5']))
        remark_text = work_order.remark if work_order.remark else ""  # Replace empty with "None"
        data = [
            ['Equipment', 'Serial Number', 'Location', 'Requester'],  # Header row
            [
                work_order.equipment.name,
                work_order.equipment.serial_number,  # Add serial number
                work_order.equipment.location,
                work_order.requester.get_full_name()
            ],
            [f"Remark: {remark_text}"]  # Use the modified remark text
        ]

        # Create table for the current work order
        table = Table(data, colWidths=[2*inch, 2*inch, 2*inch, 2*inch])  # Adjusted column widths
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),  # Header row background
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),  # Header row text color
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),  # Center align all cells
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),  # Header row font
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),  # Header row padding
            ('BACKGROUND', (0, 1), (-1, -2), colors.beige),  # Data rows background
            ('GRID', (0, 0), (-1, -1), 1, colors.black),  # Grid lines
            ('SPAN', (0, 2), (-1, 2)),  # Span the Remark row across all columns
    ('ALIGN', (0, 2), (0, 2), 'LEFT'),
    ('VALIGN', (0, 2), (0, 2), 'TOP'),  # Align to top to prevent vertical expansion
    ('LEFTPADDING', (0, 2), (0, 2), 5),  # Add some left padding
    ('RIGHTPADDING', (0, 2), (0, 2), 5),  # Add some right padding
    ('TOPPADDING', (0, 2), (0, 2), 5),  # Control top padding
    ('BOTTOMPADDING', (0, 2), (0, 2), 5),  # Control bottom padding
        ]))

        elements.append(table)

        # Add Technician and Approved By sections
        tech_approved_table_data = [
            [
                Paragraph(f"<b>Technician:</b><br/><br/>" + "<br/><br/>".join([tech.get_full_name() for tech in work_order.assigned_technicians.all()])),  # Left column
                Paragraph(f"<b>Approved By:</b><br/><br/>{work_order.approved_by.get_full_name() if work_order.approved_by else ''}")  # Right column
            ]
        ]

        tech_approved_table = Table(tech_approved_table_data, colWidths=[3*inch, 3*inch])
        tech_approved_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),  # Align text to the left
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Align text to the top
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),  # Add padding at the bottom
        ]))

        elements.append(tech_approved_table)

        # Add space between work orders
        elements.append(Paragraph("<br/><br/>", styles['Normal']))
        # elements.append(KeepTogether(elements))

    # Build the PDF
    doc.build(elements)
    return response
#------------------------------------------------------------------------------------------------------
def generate_editable_doc_work_order(work_orders, report_type, from_date, to_date):
    # Create a Word document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{report_type}_work_order_report.docx"'

    doc = Document()

    # Add main title
    title = doc.add_heading("Work Order Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.size = Pt(18)  # Make the title smaller

    if not work_orders:
        doc.add_paragraph("No work order records found for the selected date range.", style='Intense Quote')
        # Save the document to a BytesIO stream
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response.write(buffer.getvalue())
        buffer.close()
        return response

    # Add branch and date range
    branch_paragraph = doc.add_paragraph(f"Branch: {work_orders[0].branch.name}", style='Intense Quote')
    branch_paragraph.paragraph_format.space_after = Pt(0)
    date_paragraph = doc.add_paragraph(f"From {from_date} to {to_date}", style='Intense Quote')
    date_paragraph.paragraph_format.space_before = Pt(0)  # Remove space before this paragraph
    date_paragraph.paragraph_format.space_after = Pt(20)  # Add space after this paragraph

    # Iterate through each work order
    for work_order in work_orders:
        doc.add_heading(f"Work Order Date: {work_order.created_at.strftime('%Y-%m-%d')}", level=3)

        # Create table for the current work order
        table = doc.add_table(rows=2, cols=4)  # 2 rows: data + remark, 4 columns
        table.style = 'Table Grid'
        table.autofit = False

        # Set column widths
        col_widths = [2.5, 2.5, 2.5, 2]  # Widths in inches
        for i, width in enumerate(col_widths):
            table.columns[i].width = Inches(width)

        # Add table headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Equipment'
        hdr_cells[1].text = 'Serial Number'  # New column for serial number
        hdr_cells[2].text = 'Location'
        hdr_cells[3].text = 'Requester'

        # Add work order data
        row_cells = table.rows[1].cells
        row_cells[0].text = work_order.equipment.name
        row_cells[1].text = work_order.equipment.serial_number  # Add serial number
        row_cells[2].text = work_order.equipment.location
        row_cells[3].text = work_order.requester.get_full_name()
        remark_text = work_order.remark if work_order.remark else ""  # Replace empty with "None"
        # Add Remark row
        remark_row = table.add_row().cells
        remark_cell = remark_row[0]
        remark_cell.text = f"Remark: {remark_text}"
        remark_cell.merge(remark_row[1])
        remark_cell.merge(remark_row[2])
        remark_cell.merge(remark_row[3])
        
        # Format the remark cell to prevent excessive space
        for paragraph in remark_cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0  # Single spacing

        # Align "Remark:" text to the left
        for paragraph in remark_cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add space between the table and the next sections
        doc.add_paragraph()

        # Create a table for Technician and Approved By
        tech_approved_table = doc.add_table(rows=1, cols=2)
        tech_approved_table.autofit = False

        # Set column widths
        tech_approved_table.columns[0].width = Inches(3)  # Width for the "Technician" column
        tech_approved_table.columns[1].width = Inches(3)  # Width for the "Approved by" column

        # Add "Technician" to the first cell
        technician_cell = tech_approved_table.rows[0].cells[0]
        technician_cell.text = "Technician:"
        for paragraph in technician_cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True  # Make "Technician" bold

        # Add technician names below "Technician"
        for tech in work_order.assigned_technicians.all():
            technician_cell.add_paragraph(tech.get_full_name(), style='List Bullet')

        # Add "Approved by" to the second cell
        approved_by_cell = tech_approved_table.rows[0].cells[1]
        approved_by_paragraph = approved_by_cell.add_paragraph()
        approved_by_paragraph.add_run("Approved by: ").bold = True  # Make "Approved by" bold
        approved_by_paragraph.add_run(work_order.approved_by.get_full_name() if work_order.approved_by else "")

        # Add space between work orders
        doc.add_paragraph()

    # Save the document to a BytesIO stream
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response.write(buffer.getvalue())
    buffer.close()

    return response

#--------------------------------------------------------------------------------------------------------

def generate_MO_maintenance_report(from_date, to_date, export_format,report_type):
    # Fetch maintenance records for all branches
    maintenance_records = MaintenanceRecord.objects.filter(
        maintenance_type=report_type,
        status='Approved',
        datetime__date__range=[from_date, to_date],
    ).order_by('branch__name', 'datetime')

    if export_format == 'docx':
        return generate_editable_doc_all_branches(maintenance_records, report_type, from_date, to_date)
    else:
        return generate_pdf_all_branches(maintenance_records, report_type, from_date, to_date)

#--------------------------------------------------------------------------------------------------------
def generate_pdf_all_branches(maintenance_records, report_type, from_date, to_date):
    # Create a PDF document
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{report_type}_maintenance_report_all_branches.pdf"'

    doc = SimpleDocTemplate(response, pagesize=letter)
    elements = []
    doc.title = f"{report_type} Maintenance Report All Branches"

    # Add title
    styles = getSampleStyleSheet()
    title = Paragraph(f"{report_type.capitalize()} Maintenance Report (All Branches)", styles['Title'])
    elements.append(title)
    elements.append(Paragraph(f"From {from_date} to {to_date}", styles['Normal']))
    space_style = ParagraphStyle(name='Space', spaceAfter=25)
    elements.append(Paragraph("", space_style))

    # Check if maintenance_records is empty
    if not maintenance_records:
        elements.append(Paragraph("No maintenance records found for the selected date range and report type.", styles['Normal']))
        doc.build(elements)
        return response

    # Group records by branch
    records_by_branch = {}
    for record in maintenance_records:
        branch_name = record.branch.name
        if branch_name not in records_by_branch:
            records_by_branch[branch_name] = []
        records_by_branch[branch_name].append(record)

    # Iterate through each branch and add its records to the PDF
    for branch_name, records in records_by_branch.items():
        # Add branch title
        elements.append(Paragraph(f"Branch: {branch_name}", styles['Heading2']))

        # Add space between branch title and records
        space_style = ParagraphStyle(name='Space', spaceAfter=15)
        elements.append(Paragraph("", space_style))

        # Add records for the current branch
        for i, record in enumerate(records):
            if i > 0:
                elements.append(PageBreak())
            # Add header for each maintenance record
            elements.append(Paragraph(f"Maintenance Date: {record.datetime.strftime('%Y-%m-%d')}, Time: {record.datetime.strftime('%H:%M')}", styles['Heading5']))
            elements.append(Paragraph(f"Equipment: {record.equipment.name} (Serial: {record.equipment.serial_number}), Location: {record.equipment.location}", styles['Normal']))

            # Fetch all tasks associated with the maintenance task and maintenance type
            tasks = Task.objects.filter(
                task_group__maintenance_task=record.maintenance_task,
                task_group__frequency=record.maintenance_type
            )

            # Create table data
            data = [['No.', 'Inspection', 'Yes', 'No', 'Remarks']]
            for i, task in enumerate(tasks, start=1):
                # Check if the task was completed
                task_completion = TaskCompletion.objects.filter(maintenance_record=record, task=task).first()
                yes_mark = '✓' if task_completion and task_completion.is_completed else ''
                no_mark = '✓' if task_completion and not task_completion.is_completed else ''
                data.append([
                    str(i),
                    Paragraph(task.description, styles['Normal']),  # Wrap long text in a paragraph
                    yes_mark,
                    no_mark,
                    Paragraph(task_completion.remark if task_completion else '', styles['Normal'])
                ])

            # Create table with adjusted column widths
            table = Table(data, colWidths=[0.3*inch, 4*inch, 0.3*inch, 0.3*inch, 2*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Align content to the top
                ('WORDWRAP', (1, 0), (1, -1), True),  # Enable word wrap for the "Inspection" column
                ('WORDWRAP', (4, 0), (4, -1), True),  # Add this line for the "Remarks" column
                ('ALIGN', (4, 0), (4, -1), 'LEFT'),  # Optional: Left align remarks text
            ]))

            elements.append(table)

            # Add inspected by and approved by sections
            inspected_by = [tech.get_full_name() for tech in record.assigned_technicians.all()]
            approved_by = record.approved_by.get_full_name() if record.approved_by else ""

            # Create a table with one row and two columns
            table_data = [
                [
                    Paragraph(f"<b>Inspected by:</b><br/><br/>" + "<br/><br/>".join(inspected_by)),  # Left column
                    Paragraph(f"<b>Approved by:</b><br/><br/>{approved_by}")  # Right column
                ]
            ]

            # Define column widths (adjust as needed)
            col_widths = [3 * inch, 3 * inch]  # Equal width for both columns

            # Create the table
            table = Table(table_data, colWidths=col_widths)
            table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),  # Align text to the left
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Align text to the top
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),  # Add padding at the bottom
            ]))

            # Add the table to the elements
            elements.append(table)

            # Add space between records
            elements.append(Paragraph("<br/><br/>", styles['Normal']))

    # Build the PDF
    doc.build(elements)
    return response
#-------------------------------------------------------------------------------------------------------

def generate_editable_doc_all_branches(maintenance_records, report_type, from_date, to_date):
    # Create a Word document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{report_type}_maintenance_report_all_branches.docx"'

    doc = Document()

    # Add main title
    title = doc.add_heading(f"{report_type.capitalize()} Maintenance Report (All Branches)", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.size = Pt(18)  # Make the title smaller
    doc.add_paragraph(f"From {from_date} to {to_date}", style='Intense Quote')
    doc.add_paragraph()

    if not maintenance_records:
        doc.add_paragraph("No maintenance records found for the selected date range and report type.", style='Intense Quote')
        # Save the document to a BytesIO stream
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response.write(buffer.getvalue())
        buffer.close()
        return response

    # Group records by branch
    records_by_branch = {}
    for record in maintenance_records:
        branch_name = record.branch.name
        if branch_name not in records_by_branch:
            records_by_branch[branch_name] = []
        records_by_branch[branch_name].append(record)

    # Iterate through each branch and add its records to the document
    for branch_name, records in records_by_branch.items():
        # Add branch title
        doc.add_heading(f"Branch: {branch_name}", level=2)

        # Add records for the current branch
        for record in records:
            # Add header for each maintenance record
            doc.add_heading(f"Maintenance Date: {record.datetime.strftime('%Y-%m-%d')}, Time: {record.datetime.strftime('%H:%M')}", level=3)
            doc.add_paragraph(f"Equipment: {record.equipment.name} (Serial: {record.equipment.serial_number}), Location: {record.equipment.location}")

            # Create table
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            table.autofit = False

            # Set column widths
            col_widths = [0.2, 4, 0.2, 0.2, 2]  # Widths in inches
            for i, width in enumerate(col_widths):
                col = table.columns[i]
                col.width = Inches(width)

            # Add table headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'No.'
            hdr_cells[1].text = 'Inspection'
            hdr_cells[2].text = 'Yes'
            hdr_cells[3].text = 'No'
            hdr_cells[4].text = 'Remarks'

            # Fetch all tasks associated with the maintenance task and maintenance type
            tasks = Task.objects.filter(
                task_group__maintenance_task=record.maintenance_task,
                task_group__frequency=record.maintenance_type
            )

            # Add tasks to the table
            for i, task in enumerate(tasks, start=1):
                # Check if the task was completed
                task_completion = TaskCompletion.objects.filter(maintenance_record=record, task=task).first()
                yes_mark = '✓' if task_completion and task_completion.is_completed else ''
                no_mark = '✓' if task_completion and not task_completion.is_completed else ''

                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = task.description
                row_cells[2].text = yes_mark
                row_cells[3].text = no_mark
                remarks = task_completion.remark if task_completion and task_completion.remark else ''
                row_cells[4].text = str(remarks)

            # Add inspected by and approved by sections
            inspected_by = [tech.get_full_name() for tech in record.assigned_technicians.all()]
            approved_by = record.approved_by.get_full_name() if record.approved_by else ""

            # Create a table with one row and two columns
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False

            # Set column widths (adjust as needed)
            table.columns[0].width = Inches(3)  # Width for the "Inspected by" column
            table.columns[1].width = Inches(3)  # Width for the "Approved by" column

            # Add "Inspected by" to the first cell
            inspected_by_cell = table.rows[0].cells[0]
            inspected_by_cell.text = "Inspected by:"
            for paragraph in inspected_by_cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True  # Make "Inspected by" bold

            # Add technician names below "Inspected by"
            for tech in inspected_by:
                inspected_by_cell.add_paragraph(tech, style='List Bullet')

            # Add "Approved by" to the second cell
            approved_by_cell = table.rows[0].cells[1]
            approved_by_paragraph = approved_by_cell.add_paragraph()
            approved_by_paragraph.add_run("Approved by: ").bold = True  # Make "Approved by" bold
            approved_by_paragraph.add_run(approved_by)

            # Add space between records
            doc.add_paragraph()

    # Save the document to a BytesIO stream
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response.write(buffer.getvalue())
    buffer.close()

    return response
#-------------------------------------------------------------------------------------------------------

def generate_MO_work_order_report(from_date, to_date, export_format):
    # Fetch work orders for all branches
    work_orders = WorkOrder.objects.filter(
        status='Approved',
        created_at__date__range=[from_date, to_date]
    ).order_by('branch__name', 'created_at')

    if export_format == 'docx':
        return generate_editable_doc_work_order_all_branches(work_orders, from_date, to_date)
    else:
        return generate_pdf_work_order_all_branches(work_orders, from_date, to_date)
    
#---------------------------------------------------------------------------------------------------------

def generate_pdf_work_order_all_branches(work_orders, from_date, to_date):
    # Create a PDF document
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="work_order_report_all_branches.pdf"'

    doc = SimpleDocTemplate(response, pagesize=letter)
    elements = []
    doc.title = f"Work Order Report All Branches"

    # Add title
    styles = getSampleStyleSheet()
    title = Paragraph("Work Order Report (All Branches)", styles['Title'])
    elements.append(title)
    elements.append(Paragraph(f"From {from_date} to {to_date}", styles['Normal']))
    space_style = ParagraphStyle(name='Space', spaceAfter=25)
    elements.append(Paragraph("", space_style))

    # Check if work_orders is empty
    if not work_orders:
        elements.append(Paragraph("No work orders found for the selected date range.", styles['Normal']))
        doc.build(elements)
        return response

    # Group work orders by branch
    orders_by_branch = {}
    for order in work_orders:
        branch_name = order.branch.name
        if branch_name not in orders_by_branch:
            orders_by_branch[branch_name] = []
        orders_by_branch[branch_name].append(order)

    # Iterate through each branch and add its work orders to the PDF
    for branch_name, orders in orders_by_branch.items():
        # Add branch title
        elements.append(Paragraph(f"Branch: {branch_name}", styles['Heading2']))
        

        # Add space between branch title and work orders
        space_style = ParagraphStyle(name='Space', spaceAfter=30)
        elements.append(Paragraph("", space_style))

        # Add work orders for the current branch
        for order in orders:
            # Add header for each work order
            elements.append(Paragraph(f"Work Order Date: {order.created_at.strftime('%Y-%m-%d')}", styles['Heading5']))

            # Create table data
            data = [
                ['Equipment', 'Serial Number', 'Location', 'Requester'],  # Header row
                [
                    order.equipment.name,
                    order.equipment.serial_number,  # Add serial number
                    order.equipment.location,
                    order.requester.get_full_name()
                ],
                [f"Remark: {order.remark if order.remark else ''}"]  # Remark row spanning all columns
            ]

            # Create table for the current work order
            table = Table(data, colWidths=[2*inch, 2*inch, 2*inch, 2*inch])  # Adjusted column widths
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),  # Header row background
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),  # Header row text color
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),  # Center align all cells
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),  # Header row font
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),  # Header row padding
                ('BACKGROUND', (0, 1), (-1, -2), colors.beige),  # Data rows background
                ('GRID', (0, 0), (-1, -1), 1, colors.black),  # Grid lines
                ('SPAN', (0, 2), (-1, 2)),  # Span the Remark row across all columns
                ('ALIGN', (0, 2), (0, 2), 'LEFT'),  # Align "Remark:" text to the left
            ]))

            elements.append(table)

            # Add Technician and Approved By sections
            technicians = [tech.get_full_name() for tech in order.assigned_technicians.all()]
            approved_by = order.approved_by.get_full_name() if order.approved_by else ""

            # Create a table with one row and two columns
            table_data = [
                [
                    Paragraph(f"<b>Technician:</b><br/><br/>" + "<br/><br/>".join(technicians)),  # Left column
                    Paragraph(f"<b>Approved By:</b><br/><br/>{approved_by}")  # Right column
                ]
            ]

            # Define column widths (adjust as needed)
            col_widths = [3 * inch, 3 * inch]  # Equal width for both columns

            # Create the table
            table = Table(table_data, colWidths=col_widths)
            table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),  # Align text to the left
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Align text to the top
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),  # Add padding at the bottom
            ]))

            # Add the table to the elements
            elements.append(table)

            # Add space between work orders
            elements.append(Paragraph("<br/><br/>", styles['Normal']))

    # Build the PDF
    doc.build(elements)
    return response

#-----------------------------------------------------------------------------------------------------

def generate_editable_doc_work_order_all_branches(work_orders, from_date, to_date):
    # Create a Word document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="work_order_report_all_branches.docx"'

    doc = Document()

    # Add main title
    title = doc.add_heading("Work Order Report (All Branches)", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.size = Pt(18)  # Make the title smaller
    doc.add_paragraph(f"From {from_date} to {to_date}", style='Intense Quote')
    doc.add_paragraph()

    if not work_orders:
        doc.add_paragraph("No work orders found for the selected date range.", style='Intense Quote')
        # Save the document to a BytesIO stream
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response.write(buffer.getvalue())
        buffer.close()
        return response

    # Group work orders by branch
    orders_by_branch = {}
    for order in work_orders:
        branch_name = order.branch.name
        if branch_name not in orders_by_branch:
            orders_by_branch[branch_name] = []
        orders_by_branch[branch_name].append(order)

    # Iterate through each branch and add its work orders to the document
    for branch_name, orders in orders_by_branch.items():
        # Add branch title
        doc.add_heading(f"Branch: {branch_name}", level=2)

        # Add work orders for the current branch
        for order in orders:
            # Add header for each work order
            doc.add_heading(f"Work Order Date: {order.created_at.strftime('%Y-%m-%d')}", level=3)

            # Create table
            table = doc.add_table(rows=2, cols=4)  # 2 rows: data + remark, 4 columns
            table.style = 'Table Grid'
            table.autofit = False

            # Set column widths
            col_widths = [2, 2, 2, 2]  # Widths in inches
            for i, width in enumerate(col_widths):
                table.columns[i].width = Inches(width)

            # Add table headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Equipment'
            hdr_cells[1].text = 'Serial Number'  # New column for serial number
            hdr_cells[2].text = 'Location'
            hdr_cells[3].text = 'Requester'

            # Add work order data
            row_cells = table.rows[1].cells
            row_cells[0].text = order.equipment.name
            row_cells[1].text = order.equipment.serial_number  # Add serial number
            row_cells[2].text = order.equipment.location
            row_cells[3].text = order.requester.get_full_name()

            # Add Remark row
            remark_row = table.add_row().cells
            remark_cell = remark_row[0]
            remark_cell.text = f"Remark: {order.remark if order.remark else ''}"  # Add remark text on the same line

            remark_cell.merge(remark_row[1])  # Merge the first two columns
            remark_cell.merge(remark_row[2])  # Merge the next two columns
            remark_cell.merge(remark_row[3])  # Merge all columns

            

            # Align "Remark:" text to the left
            for paragraph in remark_cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Add space between the table and the next sections
            doc.add_paragraph()

            # Create a table for Technician and Approved By
            tech_approved_table = doc.add_table(rows=1, cols=2)
            tech_approved_table.autofit = False

            # Set column widths
            tech_approved_table.columns[0].width = Inches(3)  # Width for the "Technician" column
            tech_approved_table.columns[1].width = Inches(3)  # Width for the "Approved by" column

            # Add "Technician" to the first cell
            technician_cell = tech_approved_table.rows[0].cells[0]
            technician_cell.text = "Technician:"
            for paragraph in technician_cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True  # Make "Technician" bold

            # Add technician names below "Technician"
            for tech in order.assigned_technicians.all():
                technician_cell.add_paragraph(tech.get_full_name(), style='List Bullet')

            # Add "Approved by" to the second cell
            approved_by_cell = tech_approved_table.rows[0].cells[1]
            approved_by_paragraph = approved_by_cell.add_paragraph()
            approved_by_paragraph.add_run("Approved by: ").bold = True  # Make "Approved by" bold
            approved_by_paragraph.add_run(order.approved_by.get_full_name() if order.approved_by else "")

            # Add space between work orders
            doc.add_paragraph()

    # Save the document to a BytesIO stream
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response.write(buffer.getvalue())
    buffer.close()

    return response
#--------------------------------------------------------------------------------------------------------
@user_passes_test(lambda u: is_mo(u), login_url=None)
def maintenance_oversight_dashboard(request):
    user = request.user
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()


    user_role = user.userprofile.role

    # Get date range from request
    from_date = request.GET.get('from_date', (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d'))
    to_date = request.GET.get('to_date', datetime.now().strftime('%Y-%m-%d'))

    # Get selected branch from request
    selected_branch = request.GET.get('branch', 'all')
    if selected_branch != 'all':
        equipment_list = Equipment.objects.filter(branch=selected_branch).order_by('name')
    else: equipment_list = Equipment.objects.all()


    # Filter by branch if a specific branch is selected
    if selected_branch == 'all':
        branch_filter = {}
    else:
        branch_filter = {'branch': selected_branch}

    # Equipment status counts
    operational_count = Equipment.objects.filter(
        **branch_filter,
        status='operational',
        created_at__range=[from_date, to_date]
    ).count()
    non_operational_count = Equipment.objects.filter(
        **branch_filter,
        status='non_operational',
        created_at__range=[from_date, to_date]
    ).count()
    under_maintenance_count = Equipment.objects.filter(
        **branch_filter,
        status='under_maintenance',
        created_at__range=[from_date, to_date]
    ).count()

    # Maintenance by frequency counts
    daily_maintenance_count = MaintenanceRecord.objects.filter(
        **branch_filter,
        maintenance_type='daily',
        datetime__range=[from_date, to_date]
    ).count()
    weekly_maintenance_count = MaintenanceRecord.objects.filter(
        **branch_filter,
        maintenance_type='weekly',
        datetime__range=[from_date, to_date]
    ).count()
    monthly_maintenance_count = MaintenanceRecord.objects.filter(
        **branch_filter,
        maintenance_type='monthly',
        datetime__range=[from_date, to_date]
    ).count()
    quarterly_maintenance_count = MaintenanceRecord.objects.filter(
        **branch_filter,
        maintenance_type='quarterly',
        datetime__range=[from_date, to_date]
    ).count()
    biannual_maintenance_count = MaintenanceRecord.objects.filter(
        **branch_filter,
        maintenance_type='biannual',
        datetime__range=[from_date, to_date]
    ).count()
    annual_maintenance_count = MaintenanceRecord.objects.filter(
        **branch_filter,
        maintenance_type='annual',
        datetime__range=[from_date, to_date]
    ).count()

    # Work orders counts
    pending_work_orders = WorkOrder.objects.filter(
        **branch_filter,
        status='Pending',
        created_at__range=[from_date, to_date]
    ).count()
    completed_work_orders = WorkOrder.objects.filter(
        **branch_filter,
        status='Complete',
        created_at__range=[from_date, to_date]
    ).count()

    # Maintenance by month data
    maintenance_months = []
    maintenance_by_month_data = []
    for i in range(1, 13):  # January to December
        month = datetime(datetime.now().year, i, 1).strftime('%b')
        maintenance_months.append(month)
        maintenance_count = MaintenanceRecord.objects.filter(
            **branch_filter,
            datetime__month=i,
            datetime__year=datetime.now().year
        ).count()
        maintenance_by_month_data.append(maintenance_count)

    # Work orders by month data
    work_order_months = []
    pending_work_orders_by_month = []
    completed_work_orders_by_month = []
    for i in range(1, 13):  # January to December
        month = datetime(datetime.now().year, i, 1).strftime('%b')
        work_order_months.append(month)
        pending_count = WorkOrder.objects.filter(
            **branch_filter,
            created_at__month=i,
            created_at__year=datetime.now().year,
            status='Pending'
        ).count()
        completed_count = WorkOrder.objects.filter(
            **branch_filter,
            created_at__month=i,
            created_at__year=datetime.now().year,
            status='Complete'
        ).count()
        pending_work_orders_by_month.append(pending_count)
        completed_work_orders_by_month.append(completed_count)

    # Spare part usage data
    spare_part_usage = SparePartUsage.objects.filter(
        (Q(maintenance_record__branch=selected_branch) | Q(work_order__branch=selected_branch)) if selected_branch != 'all' else Q(),
        created_at__range=[from_date, to_date]
    ).values('spare_part__name').annotate(total_used=Sum('quantity_used'))

    # Extract labels and data
    spare_part_labels = [item['spare_part__name'] for item in spare_part_usage if item['spare_part__name']]
    spare_part_usage_data = [item['total_used'] for item in spare_part_usage if item['spare_part__name']]

    # Debugging
    print("Spare Part Labels:", spare_part_labels)
    print("Spare Part Usage Data:", spare_part_usage_data)

    # Fallback for empty data
    if not spare_part_labels:
        spare_part_labels = ["No Data"]
        spare_part_usage_data = [0]

    # Get all branches for the dropdown
    all_branches = Branch.objects.all()

    context = {
        'from_date': from_date,
        'to_date': to_date,
        'selected_branch': selected_branch,
        'all_branches': all_branches,
        'operational_count': operational_count,
        'non_operational_count': non_operational_count,
        'under_maintenance_count': under_maintenance_count,
        'daily_maintenance_count': daily_maintenance_count,
        'weekly_maintenance_count': weekly_maintenance_count,
        'monthly_maintenance_count': monthly_maintenance_count,
        'quarterly_maintenance_count': quarterly_maintenance_count,
        'biannual_maintenance_count': biannual_maintenance_count,
        'annual_maintenance_count': annual_maintenance_count,
        'pending_work_orders': pending_work_orders,
        'completed_work_orders': completed_work_orders,
        'maintenance_months': maintenance_months,
        'maintenance_by_month_data': maintenance_by_month_data,
        'work_order_months': work_order_months,
        'pending_work_orders_by_month': pending_work_orders_by_month,
        'completed_work_orders_by_month': completed_work_orders_by_month,
        'spare_part_labels': spare_part_labels,
        'spare_part_usage_data': spare_part_usage_data,
        'equipment_list':equipment_list,
        'active_page': 'dashboard',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }

    return render(request, 'maintenance_oversight_dashboard.html', context)
def export_maintenance_report_pdf(request):
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="maintenance_report.pdf"'

    # Create a PDF document
    doc = SimpleDocTemplate(response, pagesize=letter)
    elements = []

    # Add a title
    styles = getSampleStyleSheet()
    title = Paragraph("Maintenance Report", styles['Title'])
    elements.append(title)

    # Prepare data for the table
    maintenance_records = MaintenanceRecord.objects.all()
    data = [['Equipment', 'Maintenance Task', 'Status', 'Date', 'Tasks Completed']]
    for record in maintenance_records:
        tasks_completed = ', '.join([task.description for task in record.completed_tasks.all()])
        data.append([
            record.equipment.name,
            record.maintenance_task.equipment_type,
            record.status,
            record.datetime.strftime('%Y-%m-%d'),
            tasks_completed
        ])

    # Create a table
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))

    elements.append(table)

    # Build the PDF
    doc.build(elements)
    return response



@user_passes_test(lambda u: is_ad(u), login_url=None)
def dashboard(request):
    # Check if the user is an admin (AD)
    if request.user.userprofile.role != 'AD':
        return HttpResponse("You do not have permission to access this page.")

    # Basic counts
    total_users = User.objects.count()
    active_users = User.objects.filter(is_active=True).count()
    inactive_users = total_users - active_users
    equipment_count = Equipment.objects.count()
    
    # Equipment status
    operational_count = Equipment.objects.filter(status='Operational').count()
    non_operational_count = Equipment.objects.filter(status='Non-Operational').count()
    under_maintenance_count = Equipment.objects.filter(status='Under Maintenance').count()
    
    # User role distribution
    role_distribution = UserProfile.objects.values('role').annotate(count=models.Count('id'))
    roles = [item['role'] for item in role_distribution]
    role_counts = [item['count'] for item in role_distribution]
    
    # Users by branch with role breakdown
    branches = Branch.objects.all()
    branch_data = []
    
    for branch in branches:
        users_in_branch = UserProfile.objects.filter(branch=branch)
        branch_info = {
            'name': branch.name,
            'total': users_in_branch.count(),
            'active': users_in_branch.filter(user__is_active=True).count(),
            'inactive': users_in_branch.filter(user__is_active=False).count(),
            'roles': {}
        }
        
        # Get role distribution for this branch
        role_dist = users_in_branch.values('role').annotate(count=models.Count('id'))
        for role in role_dist:
            branch_info['roles'][role['role']] = role['count']
        
        branch_data.append(branch_info)

    context = {
        'total_users': total_users,
        'active_users': active_users,
        'inactive_users': inactive_users,
        'equipment_count': equipment_count,
        'operational_count': operational_count,
        'non_operational_count': non_operational_count,
        'under_maintenance_count': under_maintenance_count,
        'roles': roles,
        'role_counts': role_counts,
        'branch_data': branch_data,
        'active_page': 'dashboard',
    }

    return render(request, 'dashboard.html', context)

@user_passes_test(lambda u: is_cl(u), login_url=None)
def client_dashboard(request):
    user = request.user
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    # Get date range from request
    from_date = request.GET.get('from_date', (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d'))
    to_date = request.GET.get('to_date', datetime.now().strftime('%Y-%m-%d'))

    # Work order counts
    total_work_orders = WorkOrder.objects.filter(
        requester=user,
        created_at__range=[from_date, to_date]
    ).count()

    pending_work_orders = WorkOrder.objects.filter(
        requester=user,
        status='Pending',
        created_at__range=[from_date, to_date]
    ).count()

    accepted_work_orders = WorkOrder.objects.filter(
        requester=user,
        status='Accepted',
        created_at__range=[from_date, to_date]
    ).count()

    completed_work_orders = WorkOrder.objects.filter(
        requester=user,
        status='Complete',
        created_at__range=[from_date, to_date]
    ).count()

    approved_work_orders = WorkOrder.objects.filter(
        requester=user,
        status='Approved',
        created_at__range=[from_date, to_date]
    ).count()

    # Work orders by month data
    work_order_months = []
    pending_work_orders_by_month = []
    accepted_work_orders_by_month = []
    completed_work_orders_by_month = []
    approved_work_orders_by_month = []

    for i in range(1, 13):  # January to December
        month = datetime(datetime.now().year, i, 1).strftime('%b')
        work_order_months.append(month)

        pending_count = WorkOrder.objects.filter(
            requester=user,
            created_at__month=i,
            created_at__year=datetime.now().year,
            status='Pending'
        ).count()
        pending_work_orders_by_month.append(pending_count)

        accepted_count = WorkOrder.objects.filter(
            requester=user,
            created_at__month=i,
            created_at__year=datetime.now().year,
            status='Accepted'
        ).count()
        accepted_work_orders_by_month.append(accepted_count)

        completed_count = WorkOrder.objects.filter(
            requester=user,
            created_at__month=i,
            created_at__year=datetime.now().year,
            status='Complete'
        ).count()
        completed_work_orders_by_month.append(completed_count)

        approved_count = WorkOrder.objects.filter(
            requester=user,
            created_at__month=i,
            created_at__year=datetime.now().year,
            status='Approved'
        ).count()
        approved_work_orders_by_month.append(approved_count)

    context = {
        'from_date': from_date,
        'to_date': to_date,
        'total_work_orders': total_work_orders,
        'pending_work_orders': pending_work_orders,
        'accepted_work_orders': accepted_work_orders,
        'completed_work_orders': completed_work_orders,
        'approved_work_orders': approved_work_orders,
        'work_order_months': work_order_months,
        'pending_work_orders_by_month': pending_work_orders_by_month,
        'accepted_work_orders_by_month': accepted_work_orders_by_month,
        'completed_work_orders_by_month': completed_work_orders_by_month,
        'approved_work_orders_by_month': approved_work_orders_by_month,
        'active_page': 'dashboard',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }

    return render(request, 'client_dashboard.html', context)


# @sync_to_async
# def get_user_id(request):
#     close_old_connections()
#     return request.user.id

# @sync_to_async
# def get_user_id(request):
#     close_old_connections()
#     return request.user.id

# @sync_to_async
# def get_new_notifications(user_id, last_id):
#     close_old_connections()
#     return list(Notification.objects.filter(
#         user_id=user_id,
#         id__gt=last_id,
#         is_read=False
#     ).order_by('-timestamp')[:1])

# @login_required
# async def notification_stream(request):
#     async def event_generator():
#         try:
#             last_id = int(request.GET.get('last_id', 0))
#             user_id = await get_user_id(request)

#             while True:
#                 notifications = await get_new_notifications(user_id, last_id)
#                 if notifications:
#                     notification = notifications[0]
#                     last_id = notification.id
#                     yield f"data: {json.dumps({
#                         'id': notification.id,
#                         'message': notification.message,
#                         'timestamp': notification.timestamp.isoformat(),
#                         'url': notification.url
#                     })}\n\n"

#                 await asyncio.sleep(1)  # Non-blocking sleep

#         except asyncio.CancelledError:
#             # Handle client disconnection
#             pass
#         finally:
#             close_old_connections()

#     return StreamingHttpResponse(
#         event_generator(),
#         content_type='text/event-stream',
#         headers={'Cache-Control': 'no-cache'}
#     )
@user_passes_test(lambda u: is_ad(u), login_url=None)
def audit_logs(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    logs = CRUDEvent.objects.all().order_by('-datetime')
    
    # Apply filters
    event_type = request.GET.get('event_type')
    if event_type:
        # Map string values to their integer equivalents
        event_type_map = {
            'CREATE': 1,
            'UPDATE': 2,
            'DELETE': 3
        }
        if event_type in event_type_map:
            logs = logs.filter(event_type=event_type_map[event_type])
    
    user_filter = request.GET.get('user')
    if user_filter:
        logs = logs.filter(user__username__icontains=user_filter)
    
    object_type = request.GET.get('object_type')
    if object_type:
        logs = logs.filter(content_type__model__icontains=object_type)
    
    date_range = request.GET.get('date_range')
    if date_range:
        try:
            date_obj = datetime.strptime(date_range, '%Y-%m-%d').date()
            logs = logs.filter(datetime__date=date_obj)
        except ValueError:
            pass
    
    # Pagination
    paginator = Paginator(logs, 25)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'title': 'Audit Logs',
        'page_obj': page_obj,
        'active_page':'audit_log',
        'notifications':notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'audit_log.html', context)

@user_passes_test(lambda u: is_ad(u), login_url=None)
def login_events(request):
    # Get notifications (adjust according to your notification system)
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    # Get login events
    login_events = LoginEvent.objects.all().order_by('-datetime')
    
    # Get logout events (approximated from RequestEvent)
    logout_events = RequestEvent.objects.filter(
        Q(url__contains='logout') | Q(url__contains='signout'),
        datetime__gte=datetime.now() - timedelta(days=30)
    ).order_by('-datetime')
    
    # Prepare combined events with type indicator
    combined_events = []
    for event in login_events:
        combined_events.append({
            'event': event,
            'type': 'login',
            'datetime': event.datetime,
        })
    
    for event in logout_events:
        combined_events.append({
            'event': event,
            'type': 'logout',
            'datetime': event.datetime,
        })
    
    # Sort combined events by datetime
    combined_events.sort(key=lambda x: x['datetime'], reverse=True)
    
    # Apply filters
    username = request.GET.get('username')
    if username:
        combined_events = [e for e in combined_events 
                         if (e['type'] == 'login' and username.lower() in e['event'].username.lower()) or
                            (e['type'] == 'logout' and e['event'].user and username.lower() in e['event'].user.username.lower())]
    
    event_type = request.GET.get('type')
    if event_type:
        combined_events = [e for e in combined_events if e['type'] == event_type]
    
    status = request.GET.get('status')
    if status:
        combined_events = [e for e in combined_events 
                         if e['type'] == 'login' and 
                         e['event'].login_type == (1 if int(status) else 0)]
    
    date_filter = request.GET.get('date')
    if date_filter:
        try:
            filter_date = datetime.strptime(date_filter, '%Y-%m-%d').date()
            combined_events = [e for e in combined_events if e['datetime'].date() == filter_date]
        except ValueError:
            pass
    
    # Pagination
    paginator = Paginator(combined_events, 25)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'title': 'Login/Logout Events',
        'page_obj': page_obj,
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
        'active_page': 'login_event'
    }
    return render(request, 'login_events.html', context)


#-----------------------------------------------------Inventory-------------------------------------------------

@login_required
def request_spare_part_page(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    # if request.user.userprofile.role != 'TEC':
    #     messages.error(request, 'Only technicians can request spare parts.')
    #     return redirect('issue_list')
    
    user_branch = request.user.userprofile.branch
    spare_parts = SparePart.objects.filter(branch=user_branch, is_active=True)
    
    if request.method == 'POST':
        form = request.POST
        spare_part_id = form.get('spare_part')
        quantity = form.get('quantity')
        reason = form.get('reason', '')
        attachment = request.FILES.get('attachment')
        
        try:
            quantity = int(quantity)
            if quantity <= 0:
                raise ValueError("Quantity must be positive")
                
            spare_part = SparePart.objects.get(id=spare_part_id)
            
            # Check available quantity
            available = spare_part.get_available_quantity()
            if available < quantity:
                messages.error(request, f'Not enough available quantity. Only {available} available.')
                return redirect('request_spare_part_page')
            
            # Create request
            request_obj = SparePartRequest(
                technician=request.user,
                spare_part=spare_part,
                quantity_requested=quantity,
                reason=reason,
                status='Requested'
            )
            
            # Handle attachment
            if attachment:
                request_obj.attachment = attachment
            
            request_obj.save()
            
            # Notify inventory managers
            inventory_managers = User.objects.filter(
                userprofile__role='IM',
                userprofile__branch=user_branch
            )
            for manager in inventory_managers:
                Notification.objects.create(
                    user=manager,
                    type="spare_part_request",
                    message=f'New spare part request for {spare_part.name} ({quantity}) from {request.user}.',
                )
            
            messages.success(request, 'Spare part request submitted successfully!')
            return redirect('issue_list')
            
        except ValueError as e:
            messages.error(request, f'Invalid quantity: {str(e)}')
        except Exception as e:
            messages.error(request, f'Error submitting request: {str(e)}')
    
    context = {
        'spare_parts': spare_parts,
        'active_page': 'issue_list',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'request_spare_part_page.html', context)

@login_required
def request_spare_part(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    if request.method == 'POST':
        spare_part_id = request.POST.get('spare_part')
        quantity = int(request.POST.get('quantity'))
        reason = request.POST.get('reason', '')
        
        spare_part = get_object_or_404(SparePart, id=spare_part_id)
        
        # Check available quantity
        if spare_part.get_available_quantity() < quantity:
            messages.error(request, f'Not enough available quantity. Only {spare_part.get_available_quantity()} available.')
            return redirect('spare_part_list')
        
        # Create request
        SparePartRequest.objects.create(
            technician=request.user,
            spare_part=spare_part,
            quantity_requested=quantity,
            reason=reason,
            status='Requested'
        )
        
        # Notify inventory managers
        inventory_managers = User.objects.filter(userprofile__role='IM')
        for manager in inventory_managers:
            Notification.objects.create(
                user=manager,
                type="spare_part_request",
                message=f'New spare part request for {spare_part.name} (quantity : {quantity}) from {request.user}.',
            )
        
        messages.success(request, 'Spare part request submitted successfully!')
        return redirect('issue_list')
    
    return redirect('spare_part_list')

@login_required
def approve_spare_part_request(request, id):
    if not request.user.userprofile.role == 'IM':
        messages.error(request, 'You are not authorized to approve requests.')
        return redirect('issue_list')
    
    
    spare_part_request = get_object_or_404(SparePartRequest, id=id)

    if spare_part_request.status == 'Canceled':
        messages.success(request, 'The request was canceled.')
        return redirect('issue_list')
    
    if request.method == 'POST':
        action = request.POST.get('action')
        
        if action == 'approve':
            # Check available quantity again (in case it changed)
            if spare_part_request.spare_part.get_available_quantity() < spare_part_request.quantity_requested:
                messages.error(request, 'Not enough available quantity now.')
                return redirect('issue_list')
            
            spare_part_request.status = 'Approved'
            spare_part_request.inventory_manager = request.user
            spare_part_request.approval_date = timezone.now()
            spare_part_request.save()
            
            # Create transaction record
            SparePartTransaction.objects.create(
                request=spare_part_request,
                transaction_type='Approval',
                user=request.user,
                quantity=spare_part_request.quantity_requested,
                notes='Request approved'
            )
            
            # Notify technician
            Notification.objects.create(
                user=spare_part_request.technician,
                type="spare_part_request",
                message=f'Your request for {spare_part_request.spare_part.name} has been approved.',
            )
            
            messages.success(request, 'Request approved successfully!')
        
        elif action == 'reject':
            rejection_reason = request.POST.get('rejection_reason', '')
            spare_part_request.status = 'Rejected'
            spare_part_request.inventory_manager = request.user
            spare_part_request.rejection_reason = rejection_reason
            spare_part_request.save()
            
            # Create transaction record
            SparePartTransaction.objects.create(
                request=spare_part_request,
                transaction_type='Rejection',
                user=request.user,
                quantity=spare_part_request.quantity_requested,
                notes=f'Request rejected: {rejection_reason}'
            )
            
            # Notify technician
            Notification.objects.create(
                user=spare_part_request.technician,
                type="spare_part_request",
                message=f'Your request for {spare_part_request.spare_part.name} has been rejected.',
            )
            
            messages.success(request, 'Request rejected.')
    
    return redirect('issue_list')

@login_required
def issue_spare_part(request, id):
    spare_part_request = get_object_or_404(SparePartRequest, id=id, status='Approved')
    
    if request.user.userprofile.role == 'IM':
        # Inventory manager issues the part
        spare_part = spare_part_request.spare_part
        
        # Check quantity again
        if spare_part.quantity < spare_part_request.quantity_requested:
            messages.error(request, 'Not enough quantity in stock now.')
            return redirect('pending_spare_part_issues')
        
        # Update spare part quantity (reserve it)
        # spare_part.quantity -= spare_part_request.quantity_requested
        # spare_part.save()
        
        # Update request status
        spare_part_request.status = 'Issued'
        spare_part_request.issue_date = timezone.now()
        spare_part_request.save()
        
        # Create transaction record
        SparePartTransaction.objects.create(
            request=spare_part_request,
            transaction_type='Issuance',
            user=request.user,
            quantity=spare_part_request.quantity_requested,
            notes='Parts issued to technician'
        )
        check_low_spare_parts(spare_part)
        
        # Notify technician
        Notification.objects.create(
            user=spare_part_request.technician,
            type="spare_part_request",
            message=f'{spare_part_request.spare_part.name} has been issued to you.',
        )
        
        messages.success(request, 'Spare parts issued successfully!')
    
    return redirect('issue_list')



@login_required
def request_return_page(request):
    """Handle spare part return requests using get_available_quantity() method"""
    # Common context setup
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(
        user=request.user, 
        is_read=False
    ).order_by('-id').first()
    
    # Get all parts assigned to technician
    tech_parts = TechnicianSparePart.objects.filter(
        technician=request.user
    ).select_related('spare_part')
    
    # Filter parts with available quantity > 0
    tech_parts = [part for part in tech_parts if part.get_available_quantity() > 0]
    
    # Handle POST request (form submission)
    if request.method == 'POST':
        part_id = request.POST.get('spare_part')
        if not part_id:
            messages.error(request, 'Please select a spare part')
            return redirect('request_return_page')
            
        try:
            tech_part = TechnicianSparePart.objects.get(
                id=part_id,
                technician=request.user
            )
            
            return_quantity = int(request.POST.get('return_quantity', 0))
            condition = request.POST.get('condition', '')
            notes = request.POST.get('notes', '')
            
            if return_quantity <= 0:
                messages.error(request, 'Quantity must be greater than 0')
                return redirect('request_return_page')
                
            available_qty = tech_part.get_available_quantity()
            if return_quantity > available_qty:
                messages.error(request, f'You only have {available_qty} available to return')
                return redirect('request_return_page')
            
            # Create return request
            return_request = SparePartRequest.objects.create(
                technician=request.user,
                spare_part=tech_part.spare_part,
                quantity_requested=return_quantity,
                status='Return_Request',
                is_return_request=True,
                return_condition=condition,
                return_request_date=timezone.now()
            )
            
            # Handle file attachment if present
            if 'attachment' in request.FILES:
                attachment = request.FILES['attachment']
                return_request.return_attachment.save(attachment.name, attachment)
            
            # Create transaction
            SparePartTransaction.objects.create(
                request=return_request,
                transaction_type='Return_Request',
                user=request.user,
                quantity=return_quantity,
                notes=f'Return requested. Condition: {condition}'
            )
            
            # Notify inventory managers
            inventory_managers = User.objects.filter(
                userprofile__role='IM',
                userprofile__branch=request.user.userprofile.branch
            )
            for manager in inventory_managers:
                Notification.objects.create(
                    user=manager,
                    type="spare_part_request",
                    message=f'New return request for {tech_part.spare_part.name} ({return_quantity}) from {request.user}.',
                )
            
            messages.success(request, 'Return request submitted successfully!')
            return redirect('issue_list')
            
        except TechnicianSparePart.DoesNotExist:
            messages.error(request, 'Selected spare part not found')
        except Exception as e:
            messages.error(request, f'Error processing return: {str(e)}')
    
    return render(request, 'request_return_page.html', {
        'tech_parts': tech_parts,
        'active_page': 'issue_list',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    })

@login_required
def edit_return(request, id):
    """
    Edit a return request
    """
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(
        user=request.user, 
        is_read=False
    ).order_by('-id').first()
    return_request = get_object_or_404(SparePartRequest, id=id, is_return_request=True)
    
    if request.method == 'POST':
        form = request.POST
        return_request.return_condition = form.get('condition', return_request.return_condition)
        return_request.return_notes = form.get('notes', return_request.return_notes)
        
        # Handle quantity update if needed
        try:
            new_quantity = int(form.get('quantity', return_request.quantity_requested))
            if new_quantity <= 0:
                messages.error(request, 'Quantity must be positive')
            else:
                return_request.quantity_requested = new_quantity
        except ValueError:
            messages.error(request, 'Invalid quantity value')
        
        # Handle attachment update
        if 'attachment' in request.FILES:
            return_request.return_attachment = request.FILES['attachment']
        
        return_request.save()
        messages.success(request, 'Return request updated successfully')
        return redirect('issue_list')
    
    context = {
        'request_obj': return_request,
        'active_page': 'issue_list',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
        'transactions': return_request.transactions.all().order_by('-timestamp'),

    }
    return render(request, 'edit_return.html', context)

@login_required
def approve_return_request(request, id):
    return_request = get_object_or_404(SparePartRequest, id=id, is_return_request=True)
    
    if request.user.userprofile.role not in ['IM', 'MD manager']:
        messages.error(request, "Only inventory managers can approve returns")
        return redirect('issue_list')
    if return_request == 'Canceled':
        messages.success(request, "The request was canceled")
        return redirect('issue_list')
    return_request.status = 'Return_Accepted'
    return_request.return_accepted = True
    return_request.return_accepted_date = timezone.now()
    return_request.inventory_manager = request.user
    return_request.save()
    
    # Create transaction record
    SparePartTransaction.objects.create(
        request=return_request,
        transaction_type='Approval',
        user=request.user,
        quantity=return_request.quantity_requested,
        notes=f'Return approved by {request.user.get_full_name()}'
    )
    
    messages.success(request, 'Return request approved successfully')
    return redirect('issue_list')

@login_required
def reject_return_request(request, id):
    return_request = get_object_or_404(SparePartRequest, id=id, is_return_request=True)
    
    if request.user.userprofile.role not in ['IM', 'MD manager']:
        messages.error(request, "Only inventory managers can reject returns")
        return redirect('issue_list')
    
    if request.method == 'POST':
        reason = request.POST.get('rejection_reason', '')
        if not reason:
            messages.error(request, "Please provide a rejection reason")
            return redirect('edit_return', id=id)
            
        return_request.status = 'Rejected'
        return_request.rejection_reason = reason
        return_request.inventory_manager = request.user
        return_request.save()
        
        SparePartTransaction.objects.create(
            request=return_request,
            transaction_type='cancellation',
            user=request.user,
            quantity=return_request.quantity_requested,
            notes=f'Return rejected: {reason}'
        )
        
        messages.success(request, 'Return request rejected')
        return redirect('issue_list')
    
    return render(request, 'reject_return.html', {'return_request': return_request})

@login_required
def complete_return(request, id):
    return_request = get_object_or_404(SparePartRequest, id=id, is_return_request=True)
    
    if request.user.userprofile.role not in ['IM', 'MD manager']:
        messages.error(request, "Only inventory managers can complete returns")
        return redirect('issue_list')
    
    if return_request.status != 'Return_Accepted':
        messages.error(request, "Only accepted returns can be completed")
        return redirect('issue_list')
    
    # Update technician's spare part record
    tech_part = TechnicianSparePart.objects.get(
        technician=return_request.technician,
        spare_part=return_request.spare_part
    )
    tech_part.received_quantity -= return_request.quantity_requested
    tech_part.save()
    
    # Update inventory
    return_request.spare_part.quantity += return_request.quantity_requested
    return_request.spare_part.save()
    
    # Update request status
    return_request.status = 'Returned'
    return_request.return_completed_date = timezone.now()
    return_request.save()
    
    # Create transaction
    SparePartTransaction.objects.create(
        request=return_request,
        transaction_type='Return',
        user=request.user,
        quantity=return_request.quantity_requested,
        notes='Return completed and inventory updated'
    )
    
    messages.success(request, 'Return completed successfully')
    return redirect('issue_list')

@login_required
def cancel_spare_part_request(request, id):
    spare_part_request = get_object_or_404(SparePartRequest, id=id)
    notifications=  get_notifications(request.user)
    latest_notification = Notification.objects.filter(
            user=request.user, is_read=False
        ).order_by('-id').first()
    
    # Check permissions
    if not (request.user == spare_part_request.technician or 
            request.user.userprofile.role == 'IM'):
        messages.error(request, 'You are not authorized to cancel this request.')
        return redirect('spare_part_requests')
    
    # Check if request can be canceled
    if spare_part_request.status not in ['Return_Request','Accepted Returns', 'Requested', 'Approved', 'Issued']:
        messages.error(request, 'Only requested or approved requests can be canceled.')
        return redirect('issue_list')
    
    if request.method == 'POST':
        cancel_reason = request.POST.get('cancel_reason', '')
        
        # Update request status
        spare_part_request.status = 'Canceled'
        spare_part_request.cancel_reason = cancel_reason
        spare_part_request.canceled_by = request.user
        spare_part_request.cancel_date = timezone.now()
        spare_part_request.save()
        
        # Create transaction record
        SparePartTransaction.objects.create(
            request=spare_part_request,
            transaction_type='Cancellation',
            user=request.user,
            quantity=spare_part_request.quantity_requested,
            notes=f'Request canceled. Reason: {cancel_reason}'
        )
        
        # Notify relevant users
        if request.user == spare_part_request.technician:
            # Notify inventory manager if technician canceled
            if spare_part_request.inventory_manager:
                Notification.objects.create(
                    user=spare_part_request.inventory_manager,
                    type="spare_part_request",
                    message=f'Request #{spare_part_request.id} was canceled by technician.',
                )
        else:
            # Notify technician if inventory manager canceled
            Notification.objects.create(
                user=spare_part_request.technician,
                type="spare_part_request",
                message=f'Your request #{spare_part_request.id} was canceled by inventory manager.',
            )
        
        messages.success(request, 'Request canceled successfully!')
        return redirect('issue_list')
    
    # GET request - show cancel confirmation form
    context = {
        'req': spare_part_request,
        'active_page':'issue_list',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'cancel_spare_part_request.html', context)

@login_required
def issue_list(request):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    user = request.user
    user_branch = user.userprofile.branch
    role = user.userprofile.role
    
    # Base queryset with all needed relationships
    requests = SparePartRequest.objects.select_related(
        'spare_part', 
        'technician', 
        'technician__userprofile',
        'inventory_manager'
    ).order_by('-request_date')
    
    # Role-based filtering
    if role == 'TEC':
        requests = requests.filter(technician=user)
        page_title = "Inventory Request"
    elif role == 'IM':
        requests = requests.filter(spare_part__branch=user_branch)
        page_title = f"Inventory Request"
    elif role == 'MD manager':
        requests = requests.filter(
            Q(technician__userprofile__branch=user_branch) |
            Q(spare_part__branch=user_branch)
        )
        page_title = f"Inventory Request"
    else:
        requests = requests.none()
        page_title = "Inventory Request"
    
    # Status filtering
    status_filter = request.GET.get('status')
    if status_filter:
        requests = requests.filter(status=status_filter)
    
    context = {
        'requests': requests,  # Changed from 'all_requests' to match template
        'page_title': page_title,
        'active_page': 'issue_list',  # Updated to match
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
        'status_choices': SparePartRequest.STATUS_CHOICES,
        'current_status_filter': status_filter,
        'user_role': role,
    }
    
    return render(request, 'issue_list.html', context)


@login_required
def edit_request(request, id):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    spare_part_request = get_object_or_404(SparePartRequest, id=id)
    user = request.user
    user_role = user.userprofile.role
    
    # Check permissions
    if not (user == spare_part_request.technician or 
            user_role in ['IM', 'MD manager'] or
            user.is_superuser):
        messages.error(request, 'You are not authorized to edit this request.')
        return redirect('issue_list')
    
    if request.method == 'POST':
        form = request.POST
        files = request.FILES
        
        # Handle different actions based on user role and request status
        if 'update_request' in form and user == spare_part_request.technician and spare_part_request.status == 'Requested':
            # Technician updating their own pending request
            spare_part_request.quantity_requested = form.get('quantity')
            spare_part_request.reason = form.get('reason')
            if 'attachment' in files:
                spare_part_request.attachment = files['attachment']
            spare_part_request.save()
            messages.success(request, 'Request updated successfully!')
            return redirect('issue_list')
        
        elif 'approve_request' in form and user_role == 'IM' and spare_part_request.status == 'Requested':
            # IM approving request
            spare_part_request.status = 'Approved'
            spare_part_request.inventory_manager = user
            spare_part_request.approval_date = timezone.now()
            spare_part_request.save()
            
            SparePartTransaction.objects.create(
                request=spare_part_request,
                transaction_type='Approval',
                user=user,
                quantity=spare_part_request.quantity_requested,
                notes='Request approved'
            )
            
            messages.success(request, 'Request approved successfully!')
            return redirect('issue_list')
        
        # Add similar handlers for other actions (reject, issue, etc.)
    
    # Get available spare parts for dropdown (if needed)
    spare_parts = SparePart.objects.filter(branch=user.userprofile.branch, is_active=True)
    
    context = {
        'request_obj': spare_part_request,
        'spare_parts': spare_parts,
        'active_page': 'issue_list',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
        'user_role': user_role,
        'transactions': spare_part_request.transactions.all().order_by('-timestamp'),
    }
    return render(request, 'edit_request.html', context)


@login_required
def reject_spare_part_request(request, id):
    notifications = get_notifications(request.user)
    latest_notification = Notification.objects.filter(user=request.user, is_read=False).order_by('-id').first()

    # Get the request object or return 404
    spare_part_request = get_object_or_404(SparePartRequest, id=id)
    user = request.user
    
    # Permission check - only IM can reject requests
    if user.userprofile.role != 'IM':
        messages.error(request, 'You are not authorized to reject requests.')
        return redirect('issue_list')
    
    # Check if request is in a rejectable state
    if spare_part_request.status != 'Requested':
        messages.error(request, 'Only requested items can be rejected.')
        return redirect('issue_list')
    
    # Check if the request belongs to the IM's branch
    if spare_part_request.spare_part.branch != user.userprofile.branch:
        messages.error(request, 'You can only reject requests from your branch.')
        return redirect('issue_list')
    
    if request.method == 'POST':
        # Process rejection form
        rejection_reason = request.POST.get('rejection_reason', '').strip()
        
        if not rejection_reason:
            messages.error(request, 'Please provide a reason for rejection.')
            return redirect('reject_spare_part_request', id=id)
        
        try:
            # Update request status
            spare_part_request.status = 'Rejected'
            spare_part_request.inventory_manager = user
            spare_part_request.rejection_reason = rejection_reason
            spare_part_request.save()
            
            # Create transaction record
            SparePartTransaction.objects.create(
                request=spare_part_request,
                transaction_type='Rejection',
                user=user,
                quantity=spare_part_request.quantity_requested,
                notes=f'Request rejected. Reason: {rejection_reason}'
            )
            
            # Notify the technician
            Notification.objects.create(
                user=spare_part_request.technician,
                type="spare_part_request",
                message=f'Your request for {spare_part_request.spare_part.name} has been rejected. Reason: {rejection_reason}',
            )
            
            messages.success(request, 'Request has been rejected successfully.')
            return redirect('issue_list')
        
        except Exception as e:
            messages.error(request, f'An error occurred: {str(e)}')
            return redirect('reject_spare_part_request', id=id)
    
    # GET request - show rejection form
    context = {
        'request': spare_part_request,
        'active_page': 'issue_list',
        'notifications': notifications,
        'latest_notification_id': latest_notification.id if latest_notification else 0,
    }
    return render(request, 'reject_spare_part_request.html', context)

@login_required
def accept_issued_part(request, id):
    spare_part_request = get_object_or_404(SparePartRequest, id=id, status='Issued', technician=request.user)
    
    # Remove the GET handling completely
    spare_part_request.status = 'Received'
    spare_part_request.save()
    spare_part = spare_part_request.spare_part
   
    spare_part.quantity -= spare_part_request.quantity_requested
    spare_part.save()
    SparePartTransaction.objects.create(
        request=spare_part_request,
        transaction_type='Receipt',
        user=request.user,
        quantity=spare_part_request.quantity_requested,
        notes='Requester received the parts'
    )
    
    messages.success(request, 'Spare parts received successfully!')
    return redirect('issue_list')